#!/usr/bin/env python3
"""
Generate manuscript.docx and manuscript_text.md for the EDIA scoping review.
CJP / SAGE formatting: Times New Roman 12pt, double-spaced, Vancouver superscripts.
"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Output paths ────────────────────────────────────────────────────────────
BASE = os.path.dirname(os.path.abspath(__file__))
DOCX_PATH = os.path.join(BASE, "manuscript.docx")
MD_PATH   = os.path.join(BASE, "manuscript_text.md")

# ─── Helpers ─────────────────────────────────────────────────────────────────

def set_run_font(run, size_pt=12, bold=False, italic=False, superscript=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.superscript = superscript
    # Also set the theme font element to avoid Word overriding with Calibri
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"),    "Times New Roman")
    rPr.insert(0, rFonts)


def set_para_spacing(para, line_rule=WD_LINE_SPACING.DOUBLE, space_after=0, space_before=0):
    pf = para.paragraph_format
    pf.line_spacing_rule = line_rule
    pf.space_after  = Pt(space_after)
    pf.space_before = Pt(space_before)


def add_text_with_refs(para, text, font_size=12):
    """
    Parse text containing [n] or [n,m,...] style references and render them
    as superscript runs. Also handles [REF: ...] placeholders as italic superscript.
    """
    # Pattern: [digits] or [digits,digits,...] or [REF: ...]
    pattern = r'\[(\d+(?:,\s*\d+)*)\]|\[REF:[^\]]*\]'
    pos = 0
    for m in re.finditer(pattern, text):
        # Text before this match
        before = text[pos:m.start()]
        if before:
            run = para.add_run(before)
            set_run_font(run, font_size)
        # The reference itself
        ref_text = m.group(0)
        if m.group(1):
            # Numeric reference(s): superscript
            ref_nums = re.findall(r'\d+', ref_text)
            sup_text = ",".join(ref_nums)
            run = para.add_run(sup_text)
            set_run_font(run, font_size, superscript=True)
        else:
            # [REF: ...] placeholder: italic superscript
            run = para.add_run(ref_text)
            set_run_font(run, font_size, italic=True, superscript=True)
        pos = m.end()
    # Remaining text
    tail = text[pos:]
    if tail:
        run = para.add_run(tail)
        set_run_font(run, font_size)


def add_heading(doc, text, level=1, font_size=12, bold=True, after_pt=0, before_pt=6,
                italic_style=False):
    """Add a styled heading paragraph (not using Word built-in heading styles to keep TNR)."""
    para = doc.add_paragraph()
    set_para_spacing(para, line_rule=WD_LINE_SPACING.DOUBLE, space_after=after_pt, space_before=before_pt)
    run = para.add_run(text)
    set_run_font(run, font_size, bold=bold, italic=italic_style)
    return para


def add_body_para(doc, text, font_size=12, bold=False, italic=False,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line_indent=True):
    """Add a double-spaced body paragraph, parsing inline references."""
    para = doc.add_paragraph()
    set_para_spacing(para, line_rule=WD_LINE_SPACING.DOUBLE, space_after=0, space_before=0)
    para.alignment = align
    if first_line_indent:
        para.paragraph_format.first_line_indent = Cm(1.27)
    if bold or italic:
        # Simple run, no ref parsing needed for headings
        run = para.add_run(text)
        set_run_font(run, font_size, bold=bold, italic=italic)
    else:
        add_text_with_refs(para, text, font_size)
    return para


def add_blank(doc, font_size=12):
    para = doc.add_paragraph()
    set_para_spacing(para, line_rule=WD_LINE_SPACING.DOUBLE, space_after=0, space_before=0)
    run = para.add_run("")
    set_run_font(run, font_size)
    return para


def remove_all_borders(table):
    """Remove all borders from a table (will be re-added selectively)."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    # Remove existing tblBorders if present
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def set_cell_border(cell, **kwargs):
    """
    Set borders for a specific cell.
    kwargs: top, bottom, left, right; each a dict with keys val, sz, color.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, attrs in kwargs.items():
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   attrs.get("val",   "single"))
        el.set(qn("w:sz"),    str(attrs.get("sz", 4)))
        el.set(qn("w:color"), attrs.get("color", "000000"))
        el.set(qn("w:space"), "0")
        tcBorders.append(el)
    existing = tcPr.find(qn("w:tcBorders"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcBorders)


def style_table(table, num_cols):
    """Apply horizontal-only borders: top of table, bottom of header row, bottom of table."""
    # First clear all table-level borders
    remove_all_borders(table)

    rows = table.rows
    n_rows = len(rows)

    for r_idx, row in enumerate(rows):
        for c_idx, cell in enumerate(row.cells):
            borders = {}
            if r_idx == 0:
                borders["top"]    = {"val": "single", "sz": 8, "color": "000000"}
                borders["bottom"] = {"val": "single", "sz": 4, "color": "000000"}
            elif r_idx == n_rows - 1:
                borders["bottom"] = {"val": "single", "sz": 8, "color": "000000"}
            set_cell_border(cell, **borders)
            # Remove cell shading
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = tcPr.find(qn("w:shd"))
            if shd is not None:
                tcPr.remove(shd)


def set_page_number(section_para):
    """Insert a right-aligned page number field in a header paragraph."""
    section_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = section_para.add_run()
    set_run_font(run, 12)
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# ─── Content ──────────────────────────────────────────────────────────────────

TITLE = (
    "Equity, Diversity, Inclusion, and Accessibility Reporting in Canadian "
    "Mental Health Pharmacotherapy Randomized Controlled Trials (2016 to 2026): "
    "A Scoping Review Using the PROGRESS-Plus Framework"
)

ABSTRACT = [
    ("bold_label", "Objective"),
    ("body", "To characterize equity, diversity, inclusion, and accessibility (EDIA) reporting in Canadian mental health pharmacotherapy randomized controlled trials (RCTs) published between 2016 and 2026, using the PROGRESS-Plus framework, with emphasis on racialized groups and Indigenous peoples."),
    ("bold_label", "Methods"),
    ("body", "This scoping review followed PRISMA-ScR guidelines. Four databases (PubMed, Europe PMC, Scopus, OpenAlex) were searched on March 30, 2026. Eligible studies were interventional RCTs of pharmacological agents for DSM-5/ICD-10 mental disorders with at least one verified Canadian recruitment site, published 2016 to 2026 in English. Data were charted using the PROGRESS-Plus framework across 50+ variables per study. Dual screening with reconciliation was performed; data extraction used AI-assisted methods with human quality assurance (Pass 2 Cohen's κ ≥ 0.86 for all tested fields)."),
    ("bold_label", "Results"),
    ("body", "Of 10,904 records screened, 63 RCTs (N = 8,837 participants) were included. Race/ethnicity was reported in 57.1% (36/63) of trials; among reporters, 50.0% did not state a race classification framework and 61.1% used coarse granularity. Sex was reported in 95.2% (60/63) but only 6.3% (4/63) explicitly distinguished sex from gender; all published after 2021. Religion (0%), sexual orientation/gender identity (3.2%), social capital (3.2%), and intersectional analysis (0%) were rarely or never reported. Indigenous participation was documented in 11.1% (7/63), yet no trial referenced OCAP principles or reported Indigenous outcomes separately. No trial achieved full CONSORT-Equity compliance (0%); 73.0% had none. Race/ethnicity reporting was significantly associated with international collaboration (80.8% vs. 40.5%; p = 0.002) and larger sample size (p = 0.007). Among CIHR-funded trials, 31.3% (5/16) did not report race/ethnicity despite Sex- and Gender-Based Analysis Plus (SGBA+) requirements."),
    ("bold_label", "Conclusions"),
    ("body", "Substantial EDIA reporting gaps persist in Canadian mental health pharmacotherapy RCTs, particularly for race/ethnicity frameworks, sex/gender distinction, Indigenous data sovereignty, and intersectional analysis. Mandatory PROGRESS-Plus minimum reporting standards for federally funded trials could establish accountability and improve equity in Canadian clinical research."),
    ("bold_label", "Keywords"),
    ("body", "scoping review; equity; diversity; inclusion; PROGRESS-Plus; randomized controlled trial; mental health; pharmacotherapy; Canada; CONSORT-Equity"),
]

# Main manuscript sections: list of (type, text) tuples
# Types: h1, h2, body, table_placeholder
SECTIONS = [
    ("h1", "ABSTRACT"),
]
# Abstract is added separately with bold labels

INTRO_SECTIONS = [
    ("h1", "INTRODUCTION"),
    ("h2", "Background"),
    ("body", "Health equity in clinical research demands that study populations reflect the diversity of the communities trials seek to serve. Persistent underrepresentation of racialized groups, Indigenous peoples, and other equity-deserving populations in randomized controlled trials (RCTs) compromises the generalizability of evidence and perpetuates health disparities.[1,2] International efforts to address these gaps have intensified: the United States National Institutes of Health (NIH) Revitalization Act of 1993 mandated the inclusion of women and minorities in federally funded research,[3] and subsequent reporting requirements have increased race/ethnicity reporting in US trials to over 80%.[4,5]"),
    ("body", "Canada has developed its own equity-oriented policy frameworks. The Canadian Institutes of Health Research (CIHR) Sex- and Gender-Based Analysis Plus (SGBA+) policy, formalized in 2019, requires applicants to integrate sex, gender, and other identity factors into research design.[6] The Tri-Council Policy Statement: Ethical Conduct for Research Involving Humans (TCPS2), updated in 2018, includes Chapter 9 on research involving First Nations, Inuit, and Métis peoples, emphasizing community engagement and data governance.[7] The Tri-Agency (CIHR, NSERC, SSHRC) Equity, Diversity and Inclusion Action Plan, launched in 2018, further committed federal research funders to systemic change in research culture and practice.[8] Despite these policies, the extent to which Canadian clinical trials report participant diversity remains poorly characterized."),
    ("body", "The PROGRESS-Plus framework provides a structured approach to identifying equity-relevant variables in health research.[9] Originally developed for systematic reviews, it encompasses Place of residence, Race/ethnicity/culture/language, Occupation, Gender/sex, Religion, Education, Socioeconomic status (SES), and Social capital, plus additional factors including age, disability, and sexual orientation and gender identity (SOGI).[9] CONSORT-Equity extends the CONSORT reporting guidelines to promote equity-relevant reporting in RCTs.[10]"),
    ("body", "Pharmacotherapy trials warrant particular scrutiny for EDIA reporting. Known population-level variation in drug metabolism enzymes (e.g., CYP2D6, CYP2C19) is associated with ancestral population groups, influencing therapeutic efficacy and adverse effect profiles for psychotropic medications including antidepressants, antipsychotics, and mood stabilizers.[11,12] Canada's 2021 Census reported that 26.5% of the population identified as visible minorities, a share that has grown with each successive census.[13] Without adequate demographic reporting, the applicability of pharmacotherapy evidence to this increasingly diverse population cannot be assessed, and inequities in treatment safety and efficacy may go undetected."),
    ("h2", "Rationale"),
    ("body", "No scoping review has systematically assessed PROGRESS-Plus EDIA reporting in Canadian mental health pharmacotherapy RCTs. Previous reviews have focused predominantly on US contexts, where the NIH inclusion mandate has shaped reporting practices,[4,5] and have not isolated Canadian-specific patterns. Canada's unique demographic landscape (including official categories for \"visible minorities\" (Statistics Canada), constitutional recognition of Indigenous self-determination, and substantial francophone populations) requires dedicated examination."),
    ("body", "This review addresses the gap by applying the PROGRESS-Plus framework to Canadian mental health pharmacotherapy RCTs, providing the first systematic characterization of EDIA reporting practices in this specific context. The focus on pharmacotherapy, rather than all clinical research, reflects the particular relevance of population-level pharmacokinetic and pharmacogenomic variation to treatment outcomes in psychiatry."),
    ("h2", "Objectives"),
    ("body", "The primary objectives were to: (1) determine the proportion of Canadian mental health pharmacotherapy RCTs reporting participant race/ethnicity and characterize reporting completeness and granularity; (2) assess the comparability of reported race/ethnicity categories with Statistics Canada 2021 Census visible minority classifications; and (3) identify trial-level factors associated with race/ethnicity reporting."),
    ("body", "Secondary objectives were to: (1) assess how race/ethnicity data were incorporated into trial design and analysis; (2) characterize reporting of other PROGRESS-Plus variables (gender/sex, religion, education, SES, social capital, occupation, age, disability, SOGI); (3) evaluate Indigenous participation reporting and adherence to OCAP (Ownership, Control, Access, Possession) principles; and (4) compare EDIA reporting in trials published before versus after 2020, reflecting the anticipated impact of CIHR EDI policy changes."),
]

METHODS_SECTIONS = [
    ("h1", "METHODS"),
    ("h2", "Protocol and Registration"),
    ("body", "This scoping review was not prospectively registered with PROSPERO or other protocol registries. To compensate, we ensured full methodological transparency: the complete codebase including search scripts, screening reconciliation files, data extraction codebook, and analytical code is publicly available on GitHub.[14] The review followed the Joanna Briggs Institute (JBI) methodology for scoping reviews[15] and the framework proposed by Arksey and O'Malley[16] as enhanced by Levac et al.[17] Reporting adheres to the PRISMA Extension for Scoping Reviews (PRISMA-ScR; Supplementary Table S1).[18]"),
    ("h2", "Eligibility Criteria"),
    ("body", "Studies were eligible if they were: (1) interventional RCTs (parallel, crossover, or factorial design); (2) testing a pharmacological agent as the primary intervention for a DSM-5 or ICD-10 mental disorder; (3) conducted at one or more verified Canadian recruitment sites; and (4) published in English between January 1, 2016 and March 30, 2026."),
    ("body", "We excluded observational studies, non-pharmacological intervention trials, secondary or post-hoc analyses of previously published RCTs, protocols without results, systematic reviews, meta-analyses, case reports, and trials with exclusively non-Canadian recruitment sites. The eligibility criterion for Canadian involvement was refined during the review process from \"Canadian principal investigator or institutional affiliation\" to \"verified Canadian recruitment site,\" verified via ClinicalTrials.gov and published methods sections. This refinement, documented in the study repository, resulted in the exclusion of three previously included records (all with Canadian-affiliated authors but no Canadian sites; Supplementary Table S6)."),
    ("h2", "Information Sources and Search Strategy"),
    ("body", "Four electronic databases were searched on March 30, 2026: PubMed (via E-utilities API), Europe PMC (REST API), Scopus (Elsevier API), and OpenAlex (REST API). The search strategy combined three concept blocks using Boolean AND: (1) mental health disorders (MeSH terms and free-text synonyms for DSM-5/ICD-10 categories including depressive, bipolar, psychotic, anxiety, trauma, substance use, neurocognitive, neurodevelopmental, and eating disorders); (2) RCT design (Cochrane Highly Sensitive Search Strategy adaptation plus design-specific terms); and (3) Canadian context (geographic terms, institutional affiliations, and provincial/territorial names)."),
    ("body", "EDIA-specific search terms were intentionally excluded from the search strategy to avoid preferentially identifying trials that already report diversity data, which would have biased the primary outcome.[19] Full search strategies for all four databases are provided in Supplementary Table S2. Database selection was limited to sources with programmatic (API) access to ensure full reproducibility; this excluded PsycINFO and CINAHL, which is acknowledged as a limitation."),
    ("h2", "Selection of Sources of Evidence"),
    ("body", "Records were deduplicated using exact DOI match, exact PMID match, and fuzzy title matching (threshold >0.93), removing 14,497 duplicates from 54,483 raw records. After automated pre-screening filters (no abstract, outside date range, missing core terms, systematic reviews/meta-analyses), 10,904 records underwent dual independent title/abstract screening. Inter-rater agreement was 97.8% (normalized); Cohen's κ = 0.39 (fair). The low kappa reflects the extreme class imbalance inherent to scoping review screening: with a 96% exclusion base rate, agreement on the rare inclusion decisions was less well-characterized than the near-perfect exclusion agreement.[20] A false-negative validation check of approximately 5,900 excluded records identified zero missed inclusions."),
    ("body", "Of 134 records forwarded for full-text assessment (including 5 retroactively propagated after quality assurance audit), one was unobtainable, and 133 were assessed against eligibility criteria. Seventy records were excluded (35 secondary analyses, 16 no Canadian site, 10 not RCT/protocol/review, 8 non-pharmacological, 1 pre-2016 results), yielding 63 included studies."),
    ("h2", "Data Charting Process"),
    ("body", "Data were charted using a standardized extraction form based on the PROGRESS-Plus framework, encompassing over 50 variables per study organized into trial-level covariates, PROGRESS-Plus equity variables, Indigenous-specific indicators, and CONSORT-Equity compliance. The extraction codebook with decision rules and coding conventions is provided in Supplementary File S3. Per-study extracted data are provided in Supplementary Table S4."),
    ("body", "Extraction was conducted using a two-pass approach, adapting emerging methods for LLM-assisted evidence synthesis.[21] In Pass 1, large language model (LLM) AI tools extracted data from full-text PDFs under structured prompts aligned with the codebook; all extracted values were reviewed and verified by the author (human quality assurance). In Pass 2, a random sample of 20 records (32%) was independently re-extracted by separate LLM agents operating blind to Pass 1 values. Inter-rater reliability was assessed using Cohen's κ for four key fields: race reported (κ = 1.00), sex/gender distinguished (κ = 1.00), Indigenous participation (κ = 0.86), and education reported (κ = 0.89). All exceeded the pre-specified threshold of 0.80. The two discordances (2.5% cell-level rate) were resolved by consensus review. Three additional fields specified in the codebook were not independently verified in Pass 2; this partial coverage is acknowledged as a limitation. Full Pass 2 inter-rater reliability results are reported in Supplementary Table S5."),
    ("h2", "Data Items"),
    ("body", "Extracted variables included: trial-level characteristics (publication year, journal, disorder category, intervention, sample size, design, number of sites, funder type, trial registration); PROGRESS-Plus variables (place, race/ethnicity, occupation, gender/sex, religion, education, SES, social capital, age, disability, SOGI, intersectionality); Indigenous-specific indicators (participation documented, groups named, specific trial design, separate reporting, community partnership, OCAP principles, data sovereignty score 0 to 3); and CONSORT-Equity compliance (full/partial/none). Full definitions are in Supplementary File S3."),
    ("h2", "Synthesis of Results"),
    ("body", "Results were synthesized descriptively using counts and proportions with 95% Wilson score confidence intervals for each PROGRESS-Plus variable. Temporal comparison examined reporting rates in trials published pre-2020 versus post-2020 (January 2020 onward), with sensitivity analyses at the 2018 TCPS2 and 2022 CIHR EDI policy cutpoints. Race/ethnicity data were benchmarked against Statistics Canada 2021 Census visible minority categories where possible. No meta-analysis was conducted, consistent with scoping review methodology."),
    ("h2", "Patient and Public Involvement"),
    ("body", "No patients or members of the public were involved in the design, conduct, or reporting of this review, which analyzed previously published secondary data."),
]

RESULTS_SECTIONS = [
    ("h1", "RESULTS"),
    ("h2", "Selection of Sources of Evidence"),
    ("body", "The PRISMA-ScR flow diagram (Figure 1) summarizes the study selection process. Database searches identified 54,483 records (PubMed 9,964; Europe PMC 15,772; Scopus 27,983; OpenAlex 764). After removing 14,497 duplicates and applying automated pre-screening filters, 10,904 records underwent title/abstract screening. Of these, 134 were forwarded for full-text assessment, and 63 studies met all inclusion criteria. The most common full-text exclusion reasons were secondary/post-hoc analysis (n = 35, 50.0%), no verified Canadian recruitment site (n = 16, 22.9%), and non-RCT design (n = 10, 14.3%)."),
    ("figure", ("figures/figure1_prisma_flow.png", "Figure 1. PRISMA-ScR flow diagram of study selection.")),
    ("h2", "Characteristics of Included Studies"),
    ("body", "Table 1 presents the characteristics of the 63 included trials, which enrolled a combined 8,837 participants. The median sample size was 60 (range 6 to 785). The majority used parallel-group designs (76.2%; 48/63), followed by crossover designs (19.0%; 12/63). Most trials were multisite (60.3%; 38/63) and 41.3% (26/63) included international sites outside Canada."),
    ("body", "Depression was the most common disorder category (25.4%; 16/63), followed by dementia (15.9%; 10/63), bipolar/mood disorders and substance use disorders (14.3% each; 9/63), psychotic disorders (11.1%; 7/63), ADHD (7.9%; 5/63), anxiety/trauma/OCD (4.8%; 3/63), other disorders (4.8%; 3/63), and eating disorders (1.6%; 1/63). The most common funder types were foundations (30.2%; 19/63), CIHR (25.4%; 16/63), industry (20.6%; 13/63), other government agencies (11.1%; 7/63), and NIH (11.1%; 7/63)."),
    ("body", "Publication years spanned 2013 to 2026, with the majority (62/63) published between 2016 and 2026; the highest annual counts were in 2016 (n = 10) and 2025 (n = 10). One study (Malla et al., risperidone long-acting injectable for early psychosis, n = 85) was published in 2013 but was captured by database searches covering the target period and met all other eligibility criteria; its inclusion was retained as a conservative decision to avoid post-hoc exclusion, and sensitivity analysis excluding this study did not change any headline findings."),
    ("table_ref", "Table 1. Characteristics of included trials (n = 63)."),
    ("h2", "Race/Ethnicity Reporting: Primary Outcome"),
    ("figure", ("figures/figure2_progress_plus_bars.png", "Figure 2. PROGRESS-Plus variable reporting rates across 63 included Canadian mental health pharmacotherapy RCTs.")),
    ("body", "Race/ethnicity was reported in 57.1% (36/63; 95% CI 44.9 to 68.6%) of included trials (Figure 2). Among the 36 trials reporting race/ethnicity, classification frameworks were heterogeneous: 50.0% (18/36) did not state a framework, 27.8% (10/36) used US-derived categories (e.g., NIH/OMB classifications), 13.9% (5/36) used study-specific categories coded as \"Other,\" 5.6% (2/36) used self-identified categories, and only 2.8% (1/36) used Canadian Census (visible minority) categories."),
    ("body", "Race/ethnicity granularity was predominantly coarse: 61.1% (22/36) of reporters used coarse categories (3 to 5 groups with limited differentiation), 19.4% (7/36) moderate, 16.7% (6/36) binary (e.g., \"White\" vs. \"non-White\"), and only 2.8% (1/36) achieved granular reporting (≥7 well-defined categories). The median number of race/ethnicity groups reported was 4 (range 1 to 10)."),
    ("body", "Analytical use of race/ethnicity data was limited. Among all 63 trials, only 4.8% (3/63) conducted race/ethnicity subgroup analyses and 3.2% (2/63) included race/ethnicity as a covariate; the remainder used race/ethnicity descriptively only or did not collect it. Notably, no trial (0/63) acknowledged race as a social construct or discussed the distinction between race and ethnicity in its conceptual framework."),
    ("body", "Two trials collected race/ethnicity data but did not report it in their publications. One pharmacogenomics trial examining CYP2D6-guided dosing (an enzyme with known population-level frequency variation) did not collect race/ethnicity data despite the direct pharmacogenomic relevance."),
    ("body", "Among 21 trials reporting extractable overall-sample White/Caucasian proportions (excluding three trials that reported race/ethnicity by treatment arm only), the median was 84.0% (range 12.0 to 96.0%). For context, the 2021 Canadian Census reported that 26.5% of the population identified as visible minorities,[13] implying approximately 73.5% non-visible-minority. Although direct comparison is limited by differences in classification systems (most trials used US-derived or unstated categories rather than the Canadian visible minority taxonomy); the overall pattern suggests that trial populations were less diverse than the Canadian general population for many studies, though notable exceptions existed (e.g., one international trial with 12% White participants and one Vancouver-based trial with 67% White participants in a highly diverse community)."),
    ("h2", "Sex and Gender Reporting"),
    ("body", "Sex was reported in 95.2% (60/63; 59 \"Yes\" + 1 \"Partial\") of trials. Gender identity was reported in 14.3% (9/63; 6 \"Yes\" + 3 \"Partial\"). However, only 6.3% (4/63; 95% CI 2.5 to 15.2%) of trials explicitly distinguished between sex and gender in their reporting; all four were published between 2022 and 2025, suggesting an emerging but nascent awareness of this distinction."),
    ("body", "Sex-based subgroup analyses were conducted in 14.3% (9/63) of trials, with an additional 3.2% (2/63) reporting partial analyses. Despite CIHR's SGBA+ policy requiring integration of sex and gender considerations since 2019, the vast majority of trials (93.7%) continued to conflate these concepts or report only one dimension."),
    ("h2", "Other PROGRESS-Plus Variables"),
    ("body", "Table 2 presents reporting rates across all PROGRESS-Plus variables. Age was universally reported (100%; 63/63), and place of residence was nearly universal (98.4%; 62/63 including partial). Beyond these, reporting was inconsistent: education 34.9% (22/63), disability 41.3% (26/63), SES/income 28.6% (18/63), sex 95.2% (60/63), occupation 14.3% (9/63), gender 14.3% (9/63), social capital 3.2% (2/63), SOGI 3.2% (2/63), and religion 0% (0/63). No trial conducted intersectional analysis (0/63)."),
    ("body", "The two trials reporting SOGI data were both published after 2021: one documented gender dysphoria prevalence (9.6%) and one reported transgender participants (n = 2, 0.7%). Social capital was captured indirectly through measures of social support or housing stability in two trials."),
    ("table_ref", "Table 2. PROGRESS-Plus variable reporting rates across 63 included trials."),
    ("h2", "Indigenous Participation"),
    ("body", "Indigenous participation was documented in 11.1% (7/63) of trials. The most detailed reporting was in the OPTIMA trial,[22] which enrolled First Nations (16.9%) and Métis (4.8%) participants at Canadian sites funded by CIHR and the Canadian Research Initiative in Substance Misuse, the only trial to name Canadian-specific Indigenous groups. The SALOME trial[23] reported approximately 30% Aboriginal ancestry participants recruited from Vancouver's Downtown Eastside. The remaining five trials with documented Indigenous participation used non-specific or US-context labels (e.g., \"Aboriginal,\" \"American Indian/Alaska Native\") with small absolute numbers (n = 2 to 6 per study)."),
    ("body", "No trial (0/63) referenced OCAP principles (Ownership, Control, Access, Possession), described Indigenous community partnership in research governance, reported outcomes separately for Indigenous participants, or was designed specifically with or for Indigenous communities. The Indigenous data sovereignty score was 0 for all 63 studies. This finding is particularly concerning for CIHR-funded trials involving Indigenous participants, given TCPS2 Chapter 9 requirements for community engagement."),
    ("h2", "CONSORT-Equity Compliance"),
    ("body", "No trial achieved full CONSORT-Equity compliance (0/63). Partial compliance, defined as reporting equity-relevant baseline characteristics without equity-focused subgroup analyses or explicitly referencing equity frameworks, was observed in 27.0% (17/63). The majority (73.0%; 46/63) demonstrated no CONSORT-Equity compliance. This finding aligns with international data showing low awareness and adoption of equity-focused reporting guidelines in RCTs."),
    ("h2", "Funder Patterns"),
    ("body", "Table 3 displays race/ethnicity reporting by funder type. Among CIHR-funded trials, 68.8% (11/16) reported race/ethnicity, but 31.3% (5/16) did not: a notable gap given CIHR's SGBA+ policy requirements. All NIH-funded trials reported race/ethnicity (100%; 7/7), consistent with the NIH inclusion mandate. Industry-funded trials showed intermediate reporting (53.8%; 7/13), and foundation-funded trials had the lowest rate (36.8%; 7/19)."),
    ("body", "The higher race/ethnicity reporting rate in NIH-funded trials likely reflects the longstanding NIH Revitalization Act requirements, which have no direct Canadian equivalent. That nearly one-third of CIHR-funded trials omitted race/ethnicity data suggests that SGBA+ requirements have not yet been operationalized at the trial reporting level."),
    ("table_ref", "Table 3. Race/ethnicity reporting by funder type."),
    ("h2", "Trial-Level Factors Associated with Race/Ethnicity Reporting"),
    ("body", "Several trial-level characteristics were associated with race/ethnicity reporting. Trials with international sites were significantly more likely to report race/ethnicity than all-Canadian trials (80.8%, 21/26 vs. 40.5%, 15/37; Fisher's exact test p = 0.002). Multisite trials also reported race more frequently than single-site trials (68.4%, 26/38 vs. 40.0%, 10/25; p = 0.038). Trials reporting race/ethnicity had larger sample sizes (median 90, range 25 to 785) than non-reporters (median 39, range 6 to 737; Mann-Whitney U = 682, p = 0.007). By disorder category, substance use trials had the highest race reporting rate (77.8%, 7/9), while anxiety/trauma/OCD trials had the lowest (0%, 0/3), though subgroup sizes were small. These analyses were exploratory and not corrected for multiple comparisons; they should be interpreted as hypothesis-generating rather than confirmatory."),
    ("h2", "Temporal Patterns"),
    ("body", "Temporal trends showed improvement in EDIA reporting after 2020 (Supplementary Figure S1). Among pre-2020 trials (2013 to 2019; n = 26), race/ethnicity reporting was 46.2% (12/26), compared with 64.9% (24/37) in post-2020 trials (2020 to 2026), an 18.7 percentage-point increase, though this difference did not reach statistical significance (Fisher's exact test p = 0.20). The sex/gender distinction emerged exclusively in post-2022 publications (4/4 studies published 2022 to 2025), suggesting a delayed but potentially accelerating response to CIHR SGBA+ implementation."),
]

DISCUSSION_SECTIONS = [
    ("h1", "DISCUSSION"),
    ("h2", "Summary of Evidence"),
    ("body", "This scoping review provides the first systematic characterization of EDIA reporting practices in Canadian mental health pharmacotherapy RCTs using the PROGRESS-Plus framework. Among 63 included trials enrolling 8,837 participants, we found substantial and pervasive gaps across nearly all equity dimensions. Race/ethnicity, the primary outcome, was unreported in 42.9% of trials, and when reported, classification was predominantly non-standardized (50.0% unstated framework) and coarse-grained (61.1% coarse granularity). Sex and gender were conflated in 93.7% of trials. Religion, SOGI, and intersectional analysis were effectively absent (0%, 3.2%, and 0%, respectively). Indigenous data sovereignty protections were completely lacking despite documented Indigenous participation in 11.1% of trials. No trial achieved full CONSORT-Equity compliance."),
    ("body", "These findings indicate that Canada's equity-oriented research policies, including CIHR SGBA+ (2019), TCPS2 Chapter 9 (2018), and the Tri-Agency EDI Action Plan (2018), have not yet substantively translated into improved EDIA reporting in mental health pharmacotherapy trials."),
    ("h2", "Comparison with Existing Literature"),
    ("body", "Our finding of 57.1% race/ethnicity reporting in Canadian trials contrasts unfavorably with US estimates. Following the NIH Revitalization Act and subsequent reporting mandates, race/ethnicity reporting in US clinical trials has reached approximately 80%.[4,5] Turner et al. examined over 20,000 US clinical trials registered on ClinicalTrials.gov and found that 82.8% reported race/ethnicity data, with reporting rates increasing over time in response to policy requirements.[5] Oh et al. similarly found that race/ethnicity reporting increased significantly in US trials after federal mandates but remained inconsistent in non-federally funded research.[4] Canada lacks an equivalent mandatory reporting requirement, and our data suggest this policy gap translates directly to lower reporting rates."),
    ("body", "The heterogeneity of race/ethnicity classification frameworks in our sample (with 50.0% not stating any framework, 27.8% using US-derived categories, and only 2.8% using Canadian Census visible minority categories) mirrors the broader lack of standardization in international clinical research. This heterogeneity precluded meaningful benchmarking against Statistics Canada 2021 Census data, as most reporting trials used categories that were not directly comparable to the visible minority taxonomy. The predominance of US-derived frameworks (e.g., NIH/OMB categories) in Canadian-context trials raises questions about the validity of these classifications for the Canadian population and the influence of international collaborations on demographic data collection."),
    ("body", "The association between international collaboration and higher race/ethnicity reporting (80.8% vs. 40.5%; p = 0.002) likely reflects the influence of NIH inclusion requirements on multinational trials: when a US site is involved, NIH-mandated demographic collection extends to the entire trial. Similarly, larger and multisite trials may have more standardized data collection infrastructure. These findings suggest that Canadian-only trials, particularly smaller single-site studies, are where the reporting gap is most acute and where targeted policy interventions would have the greatest impact."),
    ("body", "The 0% intersectional analysis finding is consistent with international data. Intersectional approaches, which examine how multiple social identities interact to shape health outcomes, remain rare in clinical trials globally despite growing recognition of their importance.[24] The complete absence in our sample reflects both the methodological complexity of intersectional approaches in adequately powered RCTs and the lack of policy incentives to implement them in Canadian research."),
    ("body", "The complete absence of religion reporting (0/63) is striking given the well-documented relevance of religious and spiritual beliefs to mental health treatment engagement, medication adherence, and attitudes toward psychiatry.[25] In Canada's multicultural context, religious identity intersects with ethnocultural background in ways that shape access to and acceptance of pharmacotherapy. Similarly, occupation and employment status, reported in only 14.3% (9/63) of trials, represent a notable gap given the well-established bidirectional relationship between employment and mental health outcomes,[26] and the relevance of occupational status to treatment adherence, functional recovery, and socioeconomic determinants of health."),
    ("body", "The 6.3% sex/gender distinction rate, while low, represents an emerging trend post-2022 that aligns with increasing attention to CIHR's SGBA+ framework. Clayton and Tannenbaum articulated the importance of reporting both sex and gender in clinical research, noting that sex influences biological mechanisms while gender shapes health behaviors and access to care.[27] That only four trials in our sample implemented this distinction, all published between 2022 and 2025, suggests that the scientific community is beginning to respond, but progress is slow."),
    ("h2", "Canadian Policy Implications"),
    ("body", "Our findings highlight a gap between Canadian equity policy aspirations and trial-level practice. Three specific policy disconnects merit attention."),
    ("body", "First, CIHR SGBA+ requirements, introduced in 2019, mandate that applicants articulate how sex, gender, and other identity factors will be integrated into research.[6] However, 31.3% of CIHR-funded trials in our sample did not report race/ethnicity, and 93.7% of all trials failed to distinguish sex from gender. SGBA+ appears to function primarily as a grant application requirement rather than a data reporting standard; a compliance gap that could be addressed by requiring PROGRESS-Plus minimum reporting as a condition of final grant reporting.[9]"),
    ("body", "Second, TCPS2 Chapter 9 establishes ethical obligations for research involving First Nations, Inuit, and Métis peoples, including community engagement and culturally appropriate data governance.[7] Yet no trial in our sample, including those with substantial Indigenous participation, referenced OCAP principles or described Indigenous governance structures. This represents not merely a reporting gap but an ethical gap in the conduct of CIHR-funded research involving Indigenous peoples."),
    ("body", "Third, Canadian trial registries (ClinicalTrials.gov entries for Canadian trials and the now-retired ISRCTN registrations) do not require demographic data fields beyond sex. Adding structured PROGRESS-Plus fields to trial registration could normalize equity-relevant reporting from the protocol stage onward, similar to the prospective demographic collection requirements imposed by the NIH for US-registered trials."),
    ("body", "We recommend that CIHR mandate PROGRESS-Plus minimum reporting standards for all funded trials, aligning with Canada's stated commitment to equity in health research. Such a mandate would parallel the NIH model while incorporating Canada-specific dimensions including Indigenous data sovereignty. This recommendation aligns with the Truth and Reconciliation Commission's Call to Action 19, which calls for addressing the \"distinct health needs of the Métis, Inuit, and off-reserve Aboriginal peoples,\" a goal that requires, at minimum, the systematic collection and reporting of Indigenous identity data in clinical research.[28]"),
    ("h2", "Sex/Gender Distinction"),
    ("body", "The finding that 93.7% of trials conflated sex and gender underscores a fundamental implementation gap in CIHR SGBA+. Sex (biological attributes) and gender (socially constructed roles) have distinct relevance in pharmacotherapy research: sex influences pharmacokinetics and pharmacodynamics, while gender affects treatment-seeking behavior, adherence, and outcomes reporting.[27] Conflating these dimensions obscures both biological and social determinants of treatment response. In psychiatry specifically, gender-related factors including socialization, help-seeking patterns, and exposure to gender-based violence shape both the epidemiology of mental illness and treatment engagement.[29]"),
    ("body", "The temporal clustering of all four sex/gender-distinguishing studies in 2022 to 2025 suggests that awareness is increasing, possibly reflecting the lag between SGBA+ policy introduction (2019) and its integration into research practice. Notably, race/ethnicity reporting also showed temporal improvement, increasing from 46.2% (12/26) in pre-2020 trials to 64.9% (24/37) post-2020, though this difference did not reach statistical significance (Fisher's exact test p = 0.20). Whether these trends accelerate will depend on enforcement mechanisms that extend beyond grant applications to encompass peer review standards, journal reporting requirements, and trial registration norms."),
    ("h2", "Indigenous Data Sovereignty"),
    ("body", "The juxtaposition of documented Indigenous participation (11.1%) and absent data sovereignty protections (0%) represents the most ethically consequential finding of this review. Two trials illustrate the concern: the SALOME trial enrolled approximately 30% Aboriginal ancestry participants from a highly marginalized community (Vancouver's Downtown Eastside), and the OPTIMA trial included 21.7% First Nations and Métis participants across Canadian sites.[22,23] Both were CIHR-funded. Neither described OCAP principles, community governance, or separate Indigenous outcome reporting."),
    ("body", "The First Nations Information Governance Centre (FNIGC) developed OCAP principles (Ownership, Control, Access, and Possession) to assert First Nations jurisdiction over their data.[30] These principles are reflected in TCPS2 Chapter 9 and are endorsed by CIHR as part of ethical research conduct.[7] Their complete absence from trials actively enrolling Indigenous participants suggests a structural disconnect between ethics policy and research practice. The Truth and Reconciliation Commission of Canada identified this disconnect as a priority, calling for measurable goals to \"identify and close the gaps in health outcomes\" between Indigenous and non-Indigenous Canadians.[28]"),
    ("body", "This is not merely a reporting gap to be addressed by improved checklists. It reflects a deeper failure to implement Indigenous self-determination principles in the design and governance of clinical research; a failure that demands engagement beyond reporting standards, including Indigenous-led research governance models and community partnership requirements. Future trials enrolling Indigenous participants should be co-designed with Indigenous communities from inception, with data governance agreements established prior to data collection.[31]"),
    ("h2", "Limitations"),
    ("body", "Several limitations warrant consideration. First, restricting to English-language publications excluded francophone trials, potentially omitting Quebec-based research (Quebec represents approximately 23% of Canada's population). Second, our search was limited to four databases with programmatic API access, excluding PsycINFO and CINAHL; trials indexed exclusively in these databases may have been missed. Third, AI-assisted screening and extraction methods, while disclosed with reliability statistics (Cohen's κ ≥ 0.86 for extraction), represent an emerging methodology;[21] although pilot studies have demonstrated feasibility, best practices for LLM-assisted evidence synthesis are still evolving. Fourth, this review was not prospectively registered, though full transparency was ensured through public code availability. Fifth, title/abstract screening inter-rater reliability (κ = 0.39) was fair; however, this metric is difficult to interpret given the 96% exclusion base rate, and false-negative validation identified zero missed inclusions. Sixth, temporal comparison windows were unequal (pre-2020: 7 years including one 2013 study; post-2020: approximately 6 years), and the difference in race/ethnicity reporting rates (46.2% vs. 64.9%) did not reach statistical significance (p = 0.20), limiting causal inference. Seventh, Pass 2 quality control covered four of seven codebook-specified fields; the three untested fields (Indigenous data sovereignty, SES, CONSORT-Equity compliance) have lower expected variability but were not independently verified. Eighth, our findings are specific to mental health pharmacotherapy RCTs and may not generalize to other therapeutic areas or non-pharmacological trials in Canada; no comparable Canadian-specific PROGRESS-Plus review exists for benchmarking across disciplines. Ninth, the stated objective of benchmarking reported race/ethnicity categories against Statistics Canada 2021 Census visible minority classifications was substantially precluded by the heterogeneity of classification frameworks across trials, with only one of 36 reporters using Canadian Census categories, limiting this analysis to qualitative comparison rather than the quantitative benchmarking originally envisioned."),
]

CONCLUSIONS_SECTION = [
    ("h1", "CONCLUSIONS"),
    ("body", "This scoping review reveals substantial EDIA reporting gaps in Canadian mental health pharmacotherapy RCTs published between 2016 and 2026. Race/ethnicity was unreported in 42.9% of trials; when reported, classification frameworks were inconsistent and predominantly coarse. Sex and gender were conflated in the vast majority of trials. Religion, SOGI, and intersectional analysis were effectively absent from the evidence base. Most critically, Indigenous data sovereignty protections were completely lacking despite documented Indigenous participation in over one in ten trials."),
    ("body", "These findings indicate that Canada's existing equity research policies (CIHR SGBA+, TCPS2 Chapter 9, and CIHR EDI requirements) have not yet achieved their stated objectives at the level of trial reporting. We offer four specific recommendations. First, CIHR should mandate PROGRESS-Plus minimum reporting as a condition of final grant reporting for all funded trials, moving beyond the current SGBA+ framework which functions primarily at the application stage. Second, Canadian trial registries should incorporate structured PROGRESS-Plus demographic data fields, enabling prospective demographic collection from the protocol stage. Third, research involving Indigenous peoples should require documented OCAP compliance and community partnership agreements prior to ethics approval, with separate Indigenous outcome reporting as standard practice. Fourth, psychiatric journals should adopt CONSORT-Equity as a required reporting standard, with editorial checklists that prompt authors to report, or explicitly justify the omission of, equity-relevant baseline data."),
    ("body", "The finding that international collaboration and larger sample size were independently associated with race/ethnicity reporting suggests that policy interventions should specifically target the smaller, Canadian-only trials where reporting gaps are most concentrated. This pattern also underscores the influence of international research norms, particularly US NIH requirements, in driving Canadian reporting practices, even in the absence of domestic mandates."),
    ("body", "This review establishes a baseline against which future improvements in EDIA reporting can be measured. As Canada's population continues to diversify and as calls for health equity intensify, ensuring that clinical trial evidence reflects this diversity is not merely a methodological aspiration; it is an ethical obligation."),
]

REFERENCES = [
    "1. Branson RD, Davis K Jr, Butler KL. African Americans' participation in clinical research: importance, barriers, and solutions. Am J Surg. 2007;193(1):32-39. doi:10.1016/j.amjsurg.2005.11.007",
    "2. Popejoy AB, Fullerton SM. Genomics is failing on diversity. Nature. 2016;538(7624):161-164. doi:10.1038/538161a",
    "3. National Institutes of Health. NIH Guidelines on the Inclusion of Women and Minorities as Subjects in Clinical Research. Fed Regist. 1994;59(59):14508-14513.",
    "4. Oh SS, Galanter J, Thakur N, et al. Diversity in clinical and biomedical research: a promise yet to be fulfilled. PLoS Med. 2015;12(12):e1001918. doi:10.1371/journal.pmed.1001918",
    "5. Turner BE, Steinberg JR, Weeks BT, et al. Race/ethnicity reporting and representation in US clinical trials: a cohort study. Lancet Reg Health Am. 2022;11:100252. doi:10.1016/j.lana.2022.100252",
    "6. Canadian Institutes of Health Research. How to integrate sex and gender into research [Internet]. Ottawa: CIHR; 2019 [cited 2026 Apr 14]. Available from: https://cihr-irsc.gc.ca/e/50836.html",
    "7. Canadian Institutes of Health Research, Natural Sciences and Engineering Research Council of Canada, Social Sciences and Humanities Research Council. Tri-Council Policy Statement: Ethical Conduct for Research Involving Humans: TCPS 2 (2018). Ottawa; 2018.",
    "8. Canadian Institutes of Health Research, Natural Sciences and Engineering Research Council of Canada, Social Sciences and Humanities Research Council. Tri-Agency Equity, Diversity and Inclusion Action Plan 2018 to 2025 [Internet]. Ottawa: Government of Canada; 2018 [cited 2026 Apr 14]. Available from: https://www.nserc-crsng.gc.ca/InterAgency-Interorganismes/EDI-EDI/Action-Plan_Plan-dAction_eng.asp",
    "9. O'Neill J, Tabish H, Welch V, et al. Applying an equity lens to interventions: using PROGRESS ensures consideration of socially stratifying factors to illuminate inequities in health. J Clin Epidemiol. 2014;67(1):56-64. doi:10.1016/j.jclinepi.2013.08.005",
    "10. Welch VA, Norheim OF, Jull J, et al. CONSORT-Equity 2017 extension and elaboration for better reporting of health equity in randomised trials. BMJ. 2017;359:j5085. doi:10.1136/bmj.j5085",
    "11. Bradford LD. CYP2D6 allele frequency in European Caucasians, Asians, Africans and their descendants. Pharmacogenomics. 2002;3(2):229-243. doi:10.1517/14622416.3.2.229",
    "12. Koopmans AB, Braakman MH, Vinkers DJ, et al. Meta-analysis of probability estimates of worldwide variation of CYP2D6 and CYP2C19. Transl Psychiatry. 2021;11:141. doi:10.1038/s41398-020-01129-1",
    "13. Statistics Canada. Immigration and Ethnocultural Diversity Highlight Tables, 2021 Census [Internet]. Ottawa: Statistics Canada; 2022 [cited 2026 Apr 14]. Available from: https://www12.statcan.gc.ca/census-recensement/2021/dp-pd/hlt-fst/imm/index-eng.cfm",
    "14. [Author]. EDIA Reporting in Canadian Mental Health Pharmacotherapy RCTs: Scoping Review Repository [Internet]. GitHub; 2026 [cited 2026 Apr 14]. Available from: [URL available upon acceptance]",
    "15. Peters MDJ, Godfrey C, McInerney P, et al. Chapter 11: Scoping reviews. In: Aromataris E, Munn Z, editors. JBI Manual for Evidence Synthesis. JBI; 2020. doi:10.46658/JBIMES-20-12",
    "16. Arksey H, O'Malley L. Scoping studies: towards a methodological framework. Int J Soc Res Methodol. 2005;8(1):19-32. doi:10.1080/1364557032000119616",
    "17. Levac D, Colquhoun H, O'Brien KK. Scoping studies: advancing the methodology. Implement Sci. 2010;5:69. doi:10.1186/1748-5908-5-69",
    "18. Tricco AC, Lillie E, Zarin W, et al. PRISMA Extension for Scoping Reviews (PRISMA-ScR): checklist and explanation. Ann Intern Med. 2018;169(7):467-473. doi:10.7326/M18-0850",
    "19. Welch V, Petticrew M, Tugwell P, et al; PRISMA-Equity Bellagio Group. PRISMA-Equity 2012 extension: reporting guidelines for systematic reviews with a focus on health equity. PLoS Med. 2012;9(10):e1001333. doi:10.1371/journal.pmed.1001333",
    "20. Sim J, Wright CC. The kappa statistic in reliability studies: use, interpretation, and sample size requirements. Phys Ther. 2005;85(3):257-268. doi:10.1093/ptj/85.3.257",
    "21. Gartlehner G, Kahwati L, Hilscher R, et al. Data extraction for evidence synthesis using a large language model: a proof-of-concept study. Res Synth Methods. 2024;15(4):576-589. doi:10.1002/jrsm.1710",
    "22. Jutras-Aswad D, Le Foll B, Ahamad K, et al; OPTIMA Research Group. Flexible buprenorphine/naloxone model of care for reducing opioid use in individuals with prescription-type opioid use disorder: an open-label, pragmatic, noninferiority randomized controlled trial. Am J Psychiatry. 2022;179(10):726-739. doi:10.1176/appi.ajp.21090964",
    "23. Oviedo-Joekes E, Guh D, Brissette S, et al. Hydromorphone compared with diacetylmorphine for long-term opioid dependence: a randomized clinical trial. JAMA Psychiatry. 2016;73(5):447-455. doi:10.1001/jamapsychiatry.2016.0109",
    "24. Bowleg L. The problem with the phrase women and minorities; intersectionality as an important theoretical framework for public health. Am J Public Health. 2012;102(7):1267-1273. doi:10.2105/AJPH.2012.300750",
    "25. Koenig HG. Research on religion, spirituality, and mental health: a review. Can J Psychiatry. 2009;54(5):283-291. doi:10.1177/070674370905400502",
    "26. Paul KI, Moser K. Unemployment impairs mental health: meta-analyses. J Vocat Behav. 2009;74(3):264-282. doi:10.1016/j.jvb.2009.01.001",
    "27. Clayton JA, Tannenbaum C. Reporting sex, gender, or both in clinical research? JAMA. 2016;316(18):1863-1864. doi:10.1001/jama.2016.16405",
    "28. Truth and Reconciliation Commission of Canada. Truth and Reconciliation Commission of Canada: Calls to Action. Winnipeg: TRC; 2015.",
    "29. Tannenbaum C, Ellis RP, Eyssel F, et al. Sex and gender analysis improves science and engineering. Nature. 2019;575(7781):137-146. doi:10.1038/s41586-019-1657-6",
    "30. Schnarch B. Ownership, control, access, and possession (OCAP) or self-determination applied to research: a critical analysis of contemporary First Nations research and some options for First Nations communities. J Aborig Health. 2004;1(1):80-95.",
    "31. First Nations Information Governance Centre. Ownership, Control, Access and Possession (OCAP): The Path to First Nations Information Governance. Ottawa: FNIGC; 2014."
]

# Table data
TABLE1_HEADERS = ["Characteristic", "n (%)" ]
TABLE1_ROWS = [
    ["Total trials", "63"],
    ["Total participants, N", "8,837"],
    ["Median sample size (range)", "60 (6 to 785)"],
    ["", ""],
    ["Trial design", ""],
    ["  Parallel-group", "48 (76.2%)"],
    ["  Crossover", "12 (19.0%)"],
    ["  Other/factorial", "3 (4.8%)"],
    ["", ""],
    ["Number of sites", ""],
    ["  Single-site", "25 (39.7%)"],
    ["  Multisite", "38 (60.3%)"],
    ["  International sites included", "26 (41.3%)"],
    ["", ""],
    ["Disorder category", ""],
    ["  Depression", "16 (25.4%)"],
    ["  Dementia", "10 (15.9%)"],
    ["  Bipolar/mood disorder", "9 (14.3%)"],
    ["  Substance use disorder", "9 (14.3%)"],
    ["  Psychotic disorder", "7 (11.1%)"],
    ["  ADHD", "5 (7.9%)"],
    ["  Anxiety/trauma/OCD", "3 (4.8%)"],
    ["  Other", "3 (4.8%)"],
    ["  Eating disorder", "1 (1.6%)"],
    ["", ""],
    ["Funder type", ""],
    ["  Foundation", "19 (30.2%)"],
    ["  CIHR", "16 (25.4%)"],
    ["  Industry", "13 (20.6%)"],
    ["  Other government", "7 (11.1%)"],
    ["  NIH", "7 (11.1%)"],
    ["  Not reported", "1 (1.6%)"],
    ["", ""],
    ["Trial registration", ""],
    ["  Registered", "50 (79.4%)"],
    ["  Not registered/unclear", "13 (20.6%)"],
]

TABLE2_HEADERS = ["PROGRESS-Plus Variable", "Reported, n (%)", "95% CI"]
TABLE2_ROWS = [
    ["Age (mean/range)", "63 (100%)", "94.3 to 100%"],
    ["Place of residence", "62 (98.4%)", "91.5 to 99.7%"],
    ["Sex", "60 (95.2%)", "86.9 to 98.4%"],
    ["Race/ethnicity", "36 (57.1%)", "44.9 to 68.6%"],
    ["Disability", "26 (41.3%)", "30.0 to 53.6%"],
    ["Education", "22 (34.9%)", "24.3 to 47.2%"],
    ["SES/income", "18 (28.6%)", "18.9 to 40.7%"],
    ["Gender identity", "9 (14.3%)", "7.7 to 25.0%"],
    ["Occupation", "9 (14.3%)", "7.7 to 25.0%"],
    ["Sex/gender explicitly distinguished", "4 (6.3%)", "2.5 to 15.2%"],
    ["Social capital", "2 (3.2%)", "0.9 to 10.9%"],
    ["SOGI", "2 (3.2%)", "0.9 to 10.9%"],
    ["Religion", "0 (0%)", "0 to 5.7%"],
    ["Intersectional analysis", "0 (0%)", "0 to 5.7%"],
]

TABLE3_HEADERS = ["Funder", "Total trials", "Race/ethnicity reported, n (%)", "Not reported, n (%)"]
TABLE3_ROWS = [
    ["NIH", "7", "7 (100%)", "0 (0%)"],
    ["CIHR", "16", "11 (68.8%)", "5 (31.3%)"],
    ["Industry", "13", "7 (53.8%)", "6 (46.2%)"],
    ["Other government", "7", "4 (57.1%)", "3 (42.9%)"],
    ["Foundation", "19", "7 (36.8%)", "12 (63.2%)"],
    ["Not reported", "1", "0 (0%)", "1 (100%)"],
    ["All trials", "63", "36 (57.1%)", "27 (42.9%)"],
]


# ─── Word document generation ────────────────────────────────────────────────

def build_docx():
    doc = Document()

    # ── Global default style ──────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    # Paragraph format
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    style.paragraph_format.space_after  = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # ── Page setup: first section (title page) ────────────────────────────────
    section = doc.sections[0]
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)
    section.top_margin    = Cm(5)
    section.bottom_margin = Cm(5)
    # No header/footer on title page
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    # Suppress header/footer on title page by clearing them
    for para in section.header.paragraphs:
        for run in para.runs:
            run.text = ""

    # ── TITLE PAGE ────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    set_para_spacing(p, space_after=0, space_before=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TITLE PAGE: NOT FOR REVIEW")
    set_run_font(run, 12, bold=True)

    add_blank(doc)

    p = doc.add_paragraph()
    set_para_spacing(p, line_rule=WD_LINE_SPACING.DOUBLE, space_after=0, space_before=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    set_run_font(run, 12, bold=True)

    add_blank(doc)

    def title_field(label, content):
        p = doc.add_paragraph()
        set_para_spacing(p, line_rule=WD_LINE_SPACING.DOUBLE, space_after=0, space_before=0)
        r1 = p.add_run(label + ": ")
        set_run_font(r1, 12, bold=True)
        r2 = p.add_run(content)
        set_run_font(r2, 12)

    title_field("Author",
        "Gavin Thomas, BSc, University of Calgary Cumming School of Medicine, "
        "MD Program, Calgary, Alberta, Canada")
    title_field("Corresponding author",
        "Gavin Thomas, Cumming School of Medicine, University of Calgary, "
        "3330 Hospital Drive NW, Calgary, AB T2N 4N1, Canada. "
        "Email: [corresponding author email]")
    title_field("Author contributions",
        "GT conceived and designed the review, developed the search strategy, "
        "conducted database searching, performed title/abstract and full-text "
        "screening (with AI-assisted dual screening), extracted and verified all "
        "data, conducted all analyses, created all figures and tables, and drafted "
        "the manuscript. AI tools (large language models) were used as described "
        "in the Methods section; GT maintained oversight and final decision "
        "authority at all stages.")
    title_field("Conflicts of interest",
        "The author declares no conflicts of interest.")
    title_field("Funding",
        "This work received no external funding. All costs were borne by the author.")
    title_field("Acknowledgements",
        "The author thanks the open-source and open-access communities whose "
        "tools and data made this review possible, including PubMed/E-utilities, "
        "Europe PMC, OpenAlex, Crossref, Unpaywall, and the Cochrane Highly "
        "Sensitive Search Strategy for RCTs.")
    title_field("AI disclosure",
        "This scoping review used large language model (LLM) AI tools at multiple "
        "stages: (1) dual screening (title/abstract and full-text), with human "
        "reconciliation of all disagreements; (2) data extraction, with human "
        "quality assurance and Pass 2 independent re-extraction achieving "
        "Cohen's \u03ba \u2265 0.86 for all tested fields; (3) manuscript drafting "
        "assistance, with human review and revision of all content. No AI tool is "
        "listed as an author. Full methodological details are provided in the "
        "Methods section.")
    # Compute main text word count
    import re as _re
    _body_words = 0
    for sec_list in [INTRO_SECTIONS, METHODS_SECTIONS, RESULTS_SECTIONS, DISCUSSION_SECTIONS, CONCLUSIONS_SECTION]:
        for kind, text in sec_list:
            if kind == "body":
                _clean = _re.sub(r'\[\d+(?:,\s*\d+)*\]', '', text)
                _clean = _re.sub(r'\[REF:[^\]]*\]', '', _clean)
                _body_words += len(_clean.split())
    title_field("Word count",
        f"{_body_words} (main text, excluding abstract, references, tables, and figures)")

    # ── Section break (next page) before main manuscript ─────────────────────
    # The sectPr in this paragraph's pPr defines the TITLE PAGE section.
    # It must NOT carry pgNumType start=1; page restart belongs in the main section.
    p = doc.add_paragraph()
    set_para_spacing(p, space_after=0, space_before=0)
    run = p.add_run()
    set_run_font(run, 12)
    # Insert section break element (explicit next-page type for title page section)
    pPr = p._p.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")
    # Explicit next-page break type
    sectType = OxmlElement("w:type")
    sectType.set(qn("w:val"), "nextPage")
    sectPr.append(sectType)
    # Set margins for title page section
    pgMar = OxmlElement("w:pgMar")
    pgMar.set(qn("w:left"),   "1701")   # 3 cm in twips
    pgMar.set(qn("w:right"),  "1701")
    pgMar.set(qn("w:top"),    "2835")   # 5 cm in twips
    pgMar.set(qn("w:bottom"), "2835")
    sectPr.append(pgMar)
    pPr.append(sectPr)

    # ── Set up main section (the document-level section) with page numbers ────
    main_section = doc.sections[-1]  # This is the section AFTER the break
    main_section.left_margin   = Cm(3)
    main_section.right_margin  = Cm(3)
    main_section.top_margin    = Cm(5)
    main_section.bottom_margin = Cm(5)
    # Restart page numbering at 1 in the main manuscript section
    main_sectPr = main_section._sectPr
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:start"), "1")
    existing_pgNum = main_sectPr.find(qn("w:pgNumType"))
    if existing_pgNum is not None:
        main_sectPr.remove(existing_pgNum)
    main_sectPr.append(pgNumType)
    # Continuous line numbering starting at 1 in main section (title page has none)
    lnNumType = OxmlElement("w:lnNumType")
    lnNumType.set(qn("w:countBy"), "1")
    lnNumType.set(qn("w:start"), "1")
    lnNumType.set(qn("w:distance"), "360")
    lnNumType.set(qn("w:restart"), "continuous")
    existing_ln = main_sectPr.find(qn("w:lnNumType"))
    if existing_ln is not None:
        main_sectPr.remove(existing_ln)
    main_sectPr.append(lnNumType)
    # Page number in header (right-aligned)
    header = main_section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        hdr_para = header.paragraphs[0]
    else:
        hdr_para = header.add_paragraph()
    hdr_para.clear()
    set_page_number(hdr_para)

    # ── ABSTRACT ─────────────────────────────────────────────────────────────
    add_heading(doc, "ABSTRACT", before_pt=0)
    add_blank(doc)

    for (kind, text) in ABSTRACT:
        if kind == "bold_label":
            p = doc.add_paragraph()
            set_para_spacing(p, line_rule=WD_LINE_SPACING.DOUBLE, space_after=0, space_before=0)
            run = p.add_run(text)
            set_run_font(run, 12, bold=True)
        else:
            # Abstract body paragraphs: no first-line indent (they follow bold labels)
            add_body_para(doc, text, first_line_indent=False)

    # ── MAIN SECTIONS ─────────────────────────────────────────────────────────
    all_sections = (INTRO_SECTIONS + METHODS_SECTIONS +
                    RESULTS_SECTIONS + DISCUSSION_SECTIONS + CONCLUSIONS_SECTION)

    for (kind, text) in all_sections:
        if kind == "h1":
            add_blank(doc)
            add_heading(doc, text, bold=True, before_pt=0)
        elif kind == "h2":
            add_heading(doc, text, bold=False, italic_style=True, before_pt=6)
        elif kind == "table_ref":
            p = doc.add_paragraph()
            set_para_spacing(p, line_rule=WD_LINE_SPACING.DOUBLE, space_after=0, space_before=6)
            run = p.add_run(text)
            set_run_font(run, 12, italic=True)
        elif kind == "figure":
            # text is a tuple: (image_path, caption)
            img_path, caption = text
            full_path = os.path.join(BASE, img_path)
            if os.path.exists(full_path):
                p = doc.add_paragraph()
                set_para_spacing(p, line_rule=WD_LINE_SPACING.SINGLE, space_after=0, space_before=6)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(full_path, width=Cm(15))
            # Caption: centered, italic, below figure
            p = doc.add_paragraph()
            set_para_spacing(p, line_rule=WD_LINE_SPACING.DOUBLE, space_after=6, space_before=3)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(caption)
            set_run_font(run, 12, italic=True)
        else:
            add_body_para(doc, text)

    # ── TABLES ────────────────────────────────────────────────────────────────
    add_blank(doc)
    add_heading(doc, "TABLES", bold=True, before_pt=0)

    def add_table_section(title_text, headers, rows, font_size=10):
        add_blank(doc)
        p = doc.add_paragraph()
        set_para_spacing(p, line_rule=WD_LINE_SPACING.DOUBLE, space_after=0, space_before=6)
        run = p.add_run(title_text)
        set_run_font(run, 12, bold=True)  # Table titles always 12pt (body size)

        n_cols = len(headers)
        table = doc.add_table(rows=1 + len(rows), cols=n_cols)
        table.style = "Table Grid"

        # Header row
        hdr_row = table.rows[0]
        for i, h in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            set_para_spacing(p, line_rule=WD_LINE_SPACING.SINGLE, space_after=2, space_before=2)
            run = p.add_run(h)
            set_run_font(run, font_size, bold=True)

        # Data rows
        for r_idx, row_data in enumerate(rows):
            row = table.rows[r_idx + 1]
            for c_idx, val in enumerate(row_data):
                cell = row.cells[c_idx]
                cell.text = ""
                p = cell.paragraphs[0]
                set_para_spacing(p, line_rule=WD_LINE_SPACING.SINGLE, space_after=2, space_before=2)
                run = p.add_run(val)
                set_run_font(run, font_size)

        style_table(table, n_cols)

    add_table_section(
        "Table 1. Characteristics of included trials (n = 63).",
        TABLE1_HEADERS, TABLE1_ROWS
    )
    add_table_section(
        "Table 2. PROGRESS-Plus variable reporting rates across 63 included trials.",
        TABLE2_HEADERS, TABLE2_ROWS
    )
    add_table_section(
        "Table 3. Race/ethnicity reporting by funder type.",
        TABLE3_HEADERS, TABLE3_ROWS
    )

    # ── REFERENCES ────────────────────────────────────────────────────────────
    add_blank(doc)
    add_heading(doc, "REFERENCES", bold=True, before_pt=0)
    add_blank(doc)

    for ref in REFERENCES:
        p = doc.add_paragraph()
        set_para_spacing(p, line_rule=WD_LINE_SPACING.DOUBLE, space_after=0, space_before=0)
        # Hanging indent for references
        p.paragraph_format.first_line_indent = Pt(-24)
        p.paragraph_format.left_indent       = Pt(24)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(ref)
        set_run_font(run, 12)

    doc.save(DOCX_PATH)
    print(f"Saved: {DOCX_PATH}")


# ─── Markdown generation ──────────────────────────────────────────────────────

def text_to_md_refs(text):
    """Convert [n] references to <sup>n</sup> and [REF:...] to italic."""
    pattern = r'\[(\d+(?:,\s*\d+)*)\]|\[REF:[^\]]*\]'
    def replacer(m):
        if m.group(1):
            nums = re.findall(r'\d+', m.group(0))
            return f"<sup>{','.join(nums)}</sup>"
        else:
            return f"*{m.group(0)}*"
    return re.sub(pattern, replacer, text)


def table_to_md(headers, rows):
    lines = []
    lines.append("| " + " | ".join(headers) + " |")
    lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
    for row in rows:
        cells = [c.replace("|", "\\|") for c in row]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines)


def build_md():
    lines = []

    # Title
    lines.append(f"# {TITLE}")
    lines.append("")

    # Abstract
    lines.append("## ABSTRACT")
    lines.append("")
    for kind, text in ABSTRACT:
        if kind == "bold_label":
            lines.append(f"**{text}**")
        else:
            lines.append(text_to_md_refs(text))
        lines.append("")

    # Helper
    def render_sections(sec_list):
        for kind, text in sec_list:
            if kind == "h1":
                lines.append(f"## {text}")
            elif kind == "h2":
                lines.append(f"### {text}")
            elif kind == "table_ref":
                lines.append(f"*{text}*")
            elif kind == "figure":
                img_path, caption = text
                lines.append(f"*{caption}*")
                lines.append("")
                lines.append(f"[Image: {img_path}]")
            else:
                lines.append(text_to_md_refs(text))
            lines.append("")

    render_sections(INTRO_SECTIONS)
    render_sections(METHODS_SECTIONS)
    render_sections(RESULTS_SECTIONS)
    render_sections(DISCUSSION_SECTIONS)
    render_sections(CONCLUSIONS_SECTION)

    # Tables
    lines.append("## TABLES")
    lines.append("")
    lines.append("**Table 1. Characteristics of included trials (n = 63).**")
    lines.append("")
    lines.append(table_to_md(TABLE1_HEADERS, TABLE1_ROWS))
    lines.append("")
    lines.append("**Table 2. PROGRESS-Plus variable reporting rates across 63 included trials.**")
    lines.append("")
    lines.append(table_to_md(TABLE2_HEADERS, TABLE2_ROWS))
    lines.append("")
    lines.append("**Table 3. Race/ethnicity reporting by funder type.**")
    lines.append("")
    lines.append(table_to_md(TABLE3_HEADERS, TABLE3_ROWS))
    lines.append("")

    # References
    lines.append("## REFERENCES")
    lines.append("")
    for ref in REFERENCES:
        lines.append(ref)
        lines.append("")

    with open(MD_PATH, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"Saved: {MD_PATH}")


# ─── Main ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    build_docx()
    build_md()
    print("Done.")
