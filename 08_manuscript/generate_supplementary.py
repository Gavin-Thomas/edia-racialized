#!/usr/bin/env python3
"""
Generate supplementary_materials.docx for the EDIA scoping review.

Contents (all rendered in Times New Roman, CJP/SAGE-style formatting):

    Title page
    Table of contents
    Supplementary Table S1  -- PRISMA-ScR Checklist
    Supplementary Table S2  -- Full search strategies (all four databases)
    Supplementary File  S3  -- Data extraction codebook
    Supplementary Table S4  -- Per-study extracted data (key columns)
    Supplementary Table S5  -- Pass 2 inter-rater reliability
    Supplementary Table S6  -- Records excluded after criterion refinement
    Supplementary Figure S1 -- Race / Ethnicity reporting over time
    Supplementary Figure S2 -- Race / Ethnicity reporting by funder
    Supplementary Figure S3 -- Indigenous participation dot matrix
    Supplementary Figure S4 -- Sex vs. gender reporting
    Supplementary Figure S5 -- CONSORT-Equity compliance
    Supplementary Figure S6 -- PROGRESS-Plus completeness by year

All section / table / figure numbering aligns with the callouts in the main
manuscript.
"""

import os
import re
import csv

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ============================================================================
# PATHS
# ============================================================================
BASE_DIR     = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR  = os.path.dirname(BASE_DIR)
FIG_DIR      = os.path.join(BASE_DIR, 'figures')

OUT_PATH     = os.path.join(BASE_DIR, 'supplementary_materials.docx')

CHECKLIST_MD = os.path.join(BASE_DIR, 'prisma_scr_checklist.md')
CODEBOOK_MD  = os.path.join(PROJECT_DIR, '06_data_extraction',
                            'extraction_codebook.md')
EXTRACT_CSV  = os.path.join(PROJECT_DIR, '06_data_extraction',
                            'extracted_data.csv')


# ============================================================================
# DOCUMENT HELPERS
# ============================================================================
def _set_run_font(run, size=11, bold=False, italic=False, name='Times New Roman'):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)
    rPr.insert(0, rFonts)


def _set_para(para, line_rule=WD_LINE_SPACING.SINGLE,
              space_before=0, space_after=4, align=None):
    pf = para.paragraph_format
    pf.line_spacing_rule = line_rule
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if align is not None:
        para.alignment = align


def add_heading(doc, text, level=1, *, page_break_before=False):
    """level=1 top-of-section (14pt bold), level=2 sub (12pt bold italic)."""
    p = doc.add_paragraph()
    if page_break_before:
        run_break = p.add_run()
        run_break.add_break(WD_BREAK.PAGE)
    size = 14 if level == 1 else 12
    italic = (level >= 3)
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=True, italic=italic)
    _set_para(p, space_before=12 if level == 1 else 8,
              space_after=6, align=WD_ALIGN_PARAGRAPH.LEFT)
    return p


def add_body(doc, text, size=11, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size=size, italic=italic)
    _set_para(p, align=align, space_after=6)
    return p


def add_caption(doc, text, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run_font(run, size=size, italic=True)
    _set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    return p


def _cell_set_text(cell, text, *, size=10, bold=False, italic=False,
                   align=WD_ALIGN_PARAGRAPH.LEFT):
    # Replace any existing paragraph text with our styled run
    cell.text = ''
    para = cell.paragraphs[0]
    run = para.add_run(str(text))
    _set_run_font(run, size=size, bold=bold, italic=italic)
    _set_para(para, align=align, space_after=0, space_before=0)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def add_table(doc, headers, rows, *, col_widths_cm=None, header_font=10,
              body_font=9.5, zebra=True):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.autofit = False
    tbl.style = 'Table Grid'
    # Header row
    for j, h in enumerate(headers):
        _cell_set_text(tbl.rows[0].cells[j], h, size=header_font,
                       bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        # Header shading (light gray)
        tcPr = tbl.rows[0].cells[j]._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'D9D9D9')
        tcPr.append(shd)
    # Body rows
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            _cell_set_text(tbl.rows[i + 1].cells[j], v,
                           size=body_font,
                           align=WD_ALIGN_PARAGRAPH.LEFT)
            if zebra and i % 2 == 1:
                tcPr = tbl.rows[i + 1].cells[j]._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F2F2F2')
                tcPr.append(shd)
    # Column widths
    if col_widths_cm:
        for j, w in enumerate(col_widths_cm):
            for row in tbl.rows:
                row.cells[j].width = Cm(w)
    return tbl


def add_figure(doc, png_path, caption, width_cm=15.5):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_picture(png_path, width=Cm(width_cm))
    _set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_caption(doc, caption)


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ============================================================================
# MARKDOWN TABLE PARSER
# ============================================================================
def parse_md_tables(path):
    """Yield (header_list, rows) tuples for every pipe-table in a markdown file."""
    with open(path, encoding='utf-8') as f:
        lines = f.readlines()
    tables = []
    cur_header = None
    cur_rows = []
    def _split(ln):
        ln = ln.strip().strip('|')
        return [c.strip() for c in ln.split('|')]
    for ln in lines:
        if ln.strip().startswith('|') and '|' in ln.strip()[1:]:
            parts = _split(ln)
            if cur_header is None:
                cur_header = parts
                cur_rows = []
            elif all(set(p.strip()) <= set('-:') and p.strip() for p in parts):
                # separator row -- ignore
                continue
            else:
                cur_rows.append(parts)
        else:
            if cur_header is not None:
                tables.append((cur_header, cur_rows))
                cur_header = None
                cur_rows = []
    if cur_header is not None:
        tables.append((cur_header, cur_rows))
    return tables


# ============================================================================
# SECTIONS
# ============================================================================
def write_title_page(doc):
    p = doc.add_paragraph()
    run = p.add_run('Supplementary Materials')
    _set_run_font(run, size=18, bold=True)
    _set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER,
              space_before=48, space_after=18)

    p = doc.add_paragraph()
    run = p.add_run(
        'Equity, Diversity, Inclusion, and Accessibility Reporting in '
        'Canadian Mental Health Pharmacotherapy Randomized Controlled Trials '
        '(2016\u20132026): A Scoping Review Using the PROGRESS-Plus Framework'
    )
    _set_run_font(run, size=13, italic=True)
    _set_para(p, align=WD_ALIGN_PARAGRAPH.CENTER,
              space_before=0, space_after=30)

    p = doc.add_paragraph()
    run = p.add_run(
        'This document accompanies the main manuscript and contains: the '
        'PRISMA-ScR checklist, full database search strategies, the data '
        'extraction codebook, per-study extracted data, the Pass\u00A02 '
        'inter-rater reliability results, the list of records excluded after '
        'eligibility criterion refinement, and the six supplementary figures. '
        'Item numbering matches the in-text callouts in the main manuscript.'
    )
    _set_run_font(run, size=11)
    _set_para(p, align=WD_ALIGN_PARAGRAPH.LEFT,
              space_before=0, space_after=12,
              line_rule=WD_LINE_SPACING.ONE_POINT_FIVE)


def write_toc(doc):
    add_page_break(doc)
    add_heading(doc, 'Contents', level=1)
    items = [
        'Supplementary Table S1 \u2014 PRISMA-ScR Checklist',
        'Supplementary Table S2 \u2014 Full Search Strategies '
        '(PubMed, Europe PMC, Scopus, OpenAlex)',
        'Supplementary File S3 \u2014 Data Extraction Codebook',
        'Supplementary Table S4 \u2014 Per-Study Extracted Data (63 trials)',
        'Supplementary Table S5 \u2014 Pass 2 Inter-Rater Reliability',
        'Supplementary Table S6 \u2014 Records Excluded after Canadian-Site '
        'Criterion Refinement',
        'Supplementary Figure S1 \u2014 Race / Ethnicity Reporting Over Time',
        'Supplementary Figure S2 \u2014 Race / Ethnicity Reporting by Funder '
        'Type',
        'Supplementary Figure S3 \u2014 Indigenous Participation and '
        'Governance Indicators',
        'Supplementary Figure S4 \u2014 Sex vs. Gender Reporting',
        'Supplementary Figure S5 \u2014 CONSORT-Equity Compliance',
        'Supplementary Figure S6 \u2014 PROGRESS-Plus Completeness by Year '
        'and Sample Size',
    ]
    for item in items:
        p = doc.add_paragraph()
        run = p.add_run('\u2022  ' + item)
        _set_run_font(run, size=11)
        _set_para(p, space_after=2, align=WD_ALIGN_PARAGRAPH.LEFT)


def write_s1_checklist(doc):
    add_page_break(doc)
    add_heading(doc,
        'Supplementary Table S1 \u2014 PRISMA-ScR Checklist',
        level=1)
    add_body(doc,
        'Reporting guideline: PRISMA Extension for Scoping Reviews '
        '(PRISMA-ScR). Tricco AC, Lillie E, Zarin W, et al. Ann Intern Med. '
        '2018;169(7):467\u2013473. doi:10.7326/M18-0850. Item wording '
        'reproduced verbatim from Tricco et al. 2018.', italic=True)

    tables = parse_md_tables(CHECKLIST_MD)
    if not tables:
        add_body(doc, '(checklist source not found)')
        return
    header, rows = tables[0]

    # Trim to four columns for docx legibility:
    #   # | Section / Item | PRISMA-ScR item wording | Compliance status
    keep_idx = [0, 1, 2, 5]
    new_header = ['#', 'Section / Item', 'PRISMA-ScR item (Tricco et al. 2018)',
                  'Compliance status']

    body_rows = []
    for r in rows:
        if len(r) < 6:
            continue
        # Skip section-header rows like "**TITLE** | | | | |"
        if not r[2].strip() and not r[5].strip():
            # Render as a single-cell section heading row (bold)
            label = re.sub(r'\*+', '', r[1]).strip()
            if label:
                body_rows.append([f'\u25b8 {label}', '', '', ''])
            continue
        body_rows.append([r[i].strip().replace('**', '') for i in keep_idx])

    add_table(doc,
              new_header,
              body_rows,
              col_widths_cm=[1.0, 3.2, 7.5, 4.8],
              header_font=9.5,
              body_font=8.5)


def write_s2_search(doc):
    add_page_break(doc)
    add_heading(doc,
        'Supplementary Table S2 \u2014 Full Search Strategies',
        level=1)
    add_body(doc,
        'All searches executed 2026-03-30. Database selection was restricted '
        'to sources with programmatic (API) access. PsycINFO and CINAHL were '
        'not executed (acknowledged as a limitation).', italic=True)

    # Summary table
    add_heading(doc, 'Database yields', level=2)
    add_table(doc,
              ['Database', 'Platform / API', 'Date searched',
               'Records retrieved'],
              [
                ['PubMed', 'NCBI E-utilities REST API', '2026-03-30', '9,964'],
                ['Europe PMC', 'Europe PMC REST API', '2026-03-30', '15,772'],
                ['Scopus', 'Elsevier Scopus Search API', '2026-03-30', '27,983'],
                ['OpenAlex', 'OpenAlex REST API', '2026-03-30', '764'],
                ['Combined (raw)', '\u2014', '\u2014', '54,483'],
                ['After deduplication', '\u2014', '\u2014', '39,986'],
                ['After automated filtering', '\u2014', '\u2014', '10,904'],
              ],
              col_widths_cm=[4.0, 5.5, 3.5, 3.5])

    # Concept blocks (shared across databases; the concept logic is identical
    # and is what is audit-relevant for reviewers).
    add_heading(doc, 'Search concept blocks (combined with Boolean AND)',
                level=2)
    add_body(doc,
        'Concept 1 \u2014 Mental disorders: MeSH terms and free-text synonyms '
        'covering depressive disorders, bipolar and related disorders, '
        'psychotic disorders (schizophrenia and schizoaffective disorder), '
        'anxiety disorders, trauma- and stressor-related disorders '
        '(including PTSD), obsessive-compulsive and related disorders, '
        'substance-related and addictive disorders, neurocognitive disorders '
        '(including Alzheimer\u2019s dementia and related dementias), '
        'neurodevelopmental disorders (ADHD, autism spectrum disorder), and '
        'feeding and eating disorders, using DSM-5 and ICD-10 terminology.')
    add_body(doc,
        'Concept 2 \u2014 Randomized controlled trial design: Cochrane Highly '
        'Sensitive Search Strategy (sensitivity-maximizing version, 2008 '
        'revision) adapted to each platform, supplemented by design-specific '
        'terms ("parallel-group", "crossover", "factorial design", '
        '"double-blind").')
    add_body(doc,
        'Concept 3 \u2014 Canadian context: geographic terms (Canada, '
        'Canadian), provincial / territorial names, and institutional '
        'affiliation keywords (e.g., major Canadian university / hospital '
        'terms).')
    add_body(doc,
        'EDIA-specific terms were intentionally excluded from the search '
        'strategy to avoid preferentially retrieving trials that already '
        'report diversity data, which would have biased the primary '
        'outcome.', italic=True)

    # Representative query lines
    add_heading(doc, 'Representative query strings', level=2)

    pubmed_q = (
        '((("Depressive Disorder"[MeSH] OR "Bipolar Disorder"[MeSH] OR '
        '"Schizophrenia"[MeSH] OR "Anxiety Disorders"[MeSH] OR '
        '"Stress Disorders, Traumatic"[MeSH] OR '
        '"Obsessive-Compulsive Disorder"[MeSH] OR '
        '"Substance-Related Disorders"[MeSH] OR '
        '"Neurocognitive Disorders"[MeSH] OR '
        '"Attention Deficit Disorder with Hyperactivity"[MeSH] OR '
        '"Autism Spectrum Disorder"[MeSH] OR "Feeding and Eating '
        'Disorders"[MeSH] OR depressi*[tiab] OR bipolar[tiab] OR '
        'schizophren*[tiab] OR psychot*[tiab] OR anxiet*[tiab] OR PTSD[tiab] '
        'OR "obsessive compulsive"[tiab] OR "substance use"[tiab] OR '
        'dementia*[tiab] OR Alzheimer*[tiab] OR ADHD[tiab] OR autism[tiab] '
        'OR "eating disorder*"[tiab])) AND '
        '(randomized controlled trial[pt] OR controlled clinical trial[pt] '
        'OR randomized[tiab] OR placebo[tiab] OR "randomly"[tiab] OR '
        'trial[tiab]) AND (Canada[MeSH] OR Canad*[tiab] OR Ontario[tiab] '
        'OR Quebec[tiab] OR "British Columbia"[tiab] OR Alberta[tiab] OR '
        'Manitoba[tiab] OR Saskatchewan[tiab] OR "Nova Scotia"[tiab] OR '
        '"New Brunswick"[tiab] OR "Newfoundland"[tiab] OR "Prince Edward '
        'Island"[tiab] OR Yukon[tiab] OR Nunavut[tiab] OR '
        '"Northwest Territories"[tiab])) AND (2000:3000[dp])'
    )
    epmc_q = (
        '(MESH:"Depressive Disorder" OR MESH:"Bipolar Disorder" OR '
        'MESH:"Schizophrenia" OR MESH:"Anxiety Disorders" OR '
        'MESH:"Substance-Related Disorders" OR MESH:"Neurocognitive '
        'Disorders" OR depressi* OR bipolar OR schizophren* OR anxiet* OR '
        'PTSD OR "substance use" OR dementia* OR ADHD OR autism OR '
        '"eating disorder*") AND (randomized OR randomised OR RCT OR '
        'placebo OR "clinical trial") AND (Canada OR Canadian OR Ontario '
        'OR Quebec OR "British Columbia" OR Alberta OR Manitoba OR '
        'Saskatchewan OR "Nova Scotia" OR "New Brunswick" OR '
        'Newfoundland OR Yukon OR Nunavut) AND (SRC:MED OR SRC:PMC) AND '
        'PUB_YEAR:[2000 TO 3000]'
    )
    scopus_q = (
        'TITLE-ABS-KEY((depressi* OR bipolar OR schizophren* OR psychot* '
        'OR anxiet* OR PTSD OR "obsessive compulsive" OR "substance use" '
        'OR dementia* OR Alzheimer* OR ADHD OR autism OR '
        '"eating disorder*")) AND TITLE-ABS-KEY((randomized OR randomised '
        'OR RCT OR placebo OR "clinical trial")) AND '
        'AFFILCOUNTRY(Canada) AND PUBYEAR > 1999 AND '
        '(LIMIT-TO(DOCTYPE,"ar") OR LIMIT-TO(DOCTYPE,"re"))'
    )
    openalex_q = (
        '(title_and_abstract.search:"depressive disorder" OR "bipolar" OR '
        '"schizophrenia" OR "anxiety disorder" OR "PTSD" OR '
        '"substance use disorder" OR "dementia" OR "ADHD" OR '
        '"autism spectrum" OR "eating disorder") AND '
        '(title_and_abstract.search:"randomized controlled trial" OR '
        '"randomised controlled trial" OR "placebo-controlled") AND '
        '(authorships.institutions.country_code:CA) AND '
        '(publication_year:2000-2026)'
    )

    for name, q in [('PubMed', pubmed_q),
                    ('Europe PMC', epmc_q),
                    ('Scopus',    scopus_q),
                    ('OpenAlex',  openalex_q)]:
        p = doc.add_paragraph()
        run = p.add_run(name)
        _set_run_font(run, size=11, bold=True)
        _set_para(p, space_before=6, space_after=2)
        p = doc.add_paragraph()
        run = p.add_run(q)
        _set_run_font(run, size=9.5, name='Courier New')
        _set_para(p, space_after=6,
                  line_rule=WD_LINE_SPACING.SINGLE)

    add_body(doc,
        'The complete, line-numbered search strategies including concept '
        'tables, syntax notes, and reproducible Python API scripts are '
        'available at: https://github.com/gavinshadlou/edia-racialized '
        '\u2192 04_database_search/',
        italic=True)


def write_s3_codebook(doc):
    add_page_break(doc)
    add_heading(doc,
        'Supplementary File S3 \u2014 Data Extraction Codebook',
        level=1)
    add_body(doc,
        'The codebook below was used for all extraction passes. Variable '
        'definitions, decision rules, and allowable values are reproduced '
        'verbatim from the project repository '
        '(06_data_extraction/extraction_codebook.md).',
        italic=True)

    # Render the markdown file.  Handle headings, body paragraphs, bullets,
    # and pipe-tables.
    with open(CODEBOOK_MD, encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        ln = lines[i].rstrip('\n')
        stripped = ln.strip()

        if not stripped:
            i += 1
            continue

        # Skip the top-level title (already provided by our heading) + blank
        if stripped.startswith('# ') and i < 5:
            i += 1
            continue

        if stripped.startswith('### '):
            add_heading(doc, stripped[4:].strip(), level=3)
            i += 1
            continue
        if stripped.startswith('## '):
            add_heading(doc, stripped[3:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith('# '):
            add_heading(doc, stripped[2:].strip(), level=2)
            i += 1
            continue

        # Pipe-table block
        if stripped.startswith('|'):
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl_lines.append(lines[i])
                i += 1
            # parse
            rows = []
            for tl in tbl_lines:
                parts = [c.strip() for c in tl.strip().strip('|').split('|')]
                if all(set(p) <= set('-:') and p for p in parts):
                    continue
                rows.append(parts)
            if len(rows) >= 2:
                header = rows[0]
                body = rows[1:]
                # Cap very wide codebook tables to 4 columns by concatenation
                if len(header) > 4:
                    header = header[:3] + [' / '.join(header[3:])]
                    body = [r[:3] + [' | '.join(r[3:])] for r in body]
                add_table(doc, header, body,
                          header_font=9.5, body_font=8.5, zebra=True)
            continue

        # Bullet
        if stripped.startswith(('- ', '* ')):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(stripped[2:].strip())
            _set_run_font(run, size=10.5)
            _set_para(p, space_after=2,
                      line_rule=WD_LINE_SPACING.SINGLE)
            i += 1
            continue

        # Bold-only paragraph (**Label:** ...)
        # Default body paragraph
        text = stripped
        # strip surrounding ** for display (python-docx bold inline omitted)
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        add_body(doc, text, size=11)
        i += 1


def write_s4_extracted(doc):
    add_page_break(doc)
    add_heading(doc,
        'Supplementary Table S4 \u2014 Per-Study Extracted Data '
        '(n = 63 trials)',
        level=1)
    add_body(doc,
        'Key PROGRESS-Plus and Indigenous-specific extraction fields. '
        'Additional columns (70+ variables per study) are available in the '
        'project repository (06_data_extraction/extracted_data.csv). '
        'Codes: Y = Yes, P = Partial, N = No, \u2014 = field not applicable.',
        italic=True)

    with open(EXTRACT_CSV, encoding='utf-8') as f:
        reader = csv.DictReader(f)
        rows_in = list(reader)

    def _code(v):
        if v is None:
            return '\u2014'
        s = v.strip().lower()
        if s == 'yes':     return 'Y'
        if s == 'partial': return 'P'
        if s == 'no':      return 'N'
        if s == '':        return '\u2014'
        return v.strip()

    def _short_author(author):
        a = (author or '').strip()
        # Drop commas, ampersands, "et al", etc -- keep family name only
        a = re.split(r',| and | & ', a)[0].strip()
        return a[:18]

    def _consort(v):
        s = (v or '').strip().lower()
        return {'yes': 'Full', 'partial': 'Partial', 'no': 'None'}.get(s, s or '\u2014')

    def _disorder_abbr(v):
        m = {
            'Depression':           'Depr.',
            'Bipolar-Mood':         'Bipolar',
            'Psychotic':            'Psych.',
            'Anxiety-Trauma-OCD':   'Anx/OCD',
            'Substance Use':        'SUD',
            'Dementia':             'Dementia',
            'ADHD':                 'ADHD',
            'Eating Disorder':      'Eating',
            'Other':                'Other',
        }
        return m.get((v or '').strip(), v or '')

    headers = [
        'Rec #', 'First author', 'Year', 'Disorder', 'N',
        'Race', 'Sex', 'Gender', 'S/G dist.', 'Indig.', 'CONSORT-Eq'
    ]

    # Sort by record number (numeric)
    def _rec_key(r):
        try:
            return int(r.get('record_number', '').strip())
        except Exception:
            return 9_999_999
    rows_in.sort(key=_rec_key)

    body = []
    for r in rows_in:
        body.append([
            r.get('record_number', '').strip(),
            _short_author(r.get('first_author', '')),
            r.get('year', '').strip(),
            _disorder_abbr(r.get('disorder_category', '')),
            r.get('sample_size', '').strip(),
            _code(r.get('race_reported')),
            _code(r.get('sex_reported')),
            _code(r.get('gender_reported')),
            _code(r.get('sex_gender_distinguished')),
            _code(r.get('indigenous_participation')),
            _consort(r.get('consort_equity_compliant')),
        ])

    add_table(doc, headers, body,
              col_widths_cm=[1.1, 3.0, 1.0, 1.7, 1.1, 1.0, 1.0, 1.1, 1.2, 1.1, 1.7],
              header_font=9, body_font=8.5)


def write_s5_irr(doc):
    add_page_break(doc)
    add_heading(doc,
        'Supplementary Table S5 \u2014 Pass 2 Inter-Rater Reliability '
        '(Independent Re-Extraction, n = 20 trials / 32%)',
        level=1)
    add_body(doc,
        'A random sample of 20 / 63 records (32%) was independently '
        're-extracted by separate LLM agents operating blind to Pass\u00A01 '
        'values. Cohen\u2019s \u03BA is reported for dichotomous / categorical '
        'fields. All four fields exceeded the pre-specified reliability '
        'threshold of \u03BA \u2265 0.80. Two cell-level discordances (2.5%) '
        'were resolved by consensus review.', italic=True)

    add_table(doc,
              ['Field',
               'n agreed',
               'n total',
               'Cohen\u2019s \u03BA',
               'Interpretation'],
              [
                ['race_reported',
                 '20', '20', '1.00', 'Perfect agreement'],
                ['sex_gender_distinguished',
                 '20', '20', '1.00', 'Perfect agreement'],
                ['indigenous_participation',
                 '19', '20', '0.86', 'Almost perfect (Landis & Koch 1977)'],
                ['education_reported',
                 '19', '20', '0.89', 'Almost perfect (Landis & Koch 1977)'],
              ],
              col_widths_cm=[5.0, 2.0, 2.0, 2.5, 5.0])

    add_body(doc,
        'Pre-specified reliability threshold: \u03BA \u2265 0.80 for all '
        'fields. Fields not independently re-extracted in Pass 2 '
        '(indigenous_data_sovereignty, ses_reported, '
        'consort_equity_compliant) are acknowledged as a limitation in the '
        'main manuscript.', italic=True)


def write_s6_excluded(doc):
    add_page_break(doc)
    add_heading(doc,
        'Supplementary Table S6 \u2014 Records Excluded after Canadian-Site '
        'Criterion Refinement (2026-04-11)',
        level=1)
    add_body(doc,
        'The eligibility criterion was simplified from '
        '\u201CCanadian principal investigator or institutional affiliation\u201D '
        'to \u201Cverified Canadian recruitment site,\u201D verified via '
        'ClinicalTrials.gov and published methods. The three records below '
        'were retroactively reclassified as EXCLUDED at full-text review. '
        'Full PDFs remain on disk in the project repository as audit-trail '
        'evidence.', italic=True)

    add_table(doc,
              ['Record #', 'PMID', 'Study',
               'Canadian author affiliation',
               'Actual recruitment sites'],
              [
                ['20', '37227402',
                 'McIntyre et al. 2023 \u2014 VIVRE: vortioxetine vs. '
                 'desvenlafaxine for MDD',
                 'McIntyre (University of Toronto) \u2014 lead author, not a '
                 'site PI',
                 '80 sites / 12 countries (Russia, Argentina, Ukraine, '
                 'and others); no Canadian site'],
                ['42', '39144112',
                 'Husain et al. 2024 \u2014 Sodium benzoate + N-acetylcysteine '
                 'feasibility study in schizophrenia',
                 'CAMH authors \u2014 design contribution only, no '
                 'recruitment',
                 '5 sites, all in Pakistan'],
                ['95', '28044255',
                 'Rutrick et al. 2017 \u2014 Mavoglurant for OCD',
                 'Gomez-Mancilla (McGill secondary, Novartis Basel '
                 'primary)',
                 '15 sites in Bulgaria / Germany / USA / Czechia / '
                 'Switzerland'],
              ],
              col_widths_cm=[1.4, 2.0, 4.5, 4.0, 4.8])


def write_figs(doc):
    fig_list = [
        ('figure3_race_over_time.png',
         'Supplementary Figure S1. Race / ethnicity reporting across five '
         'publication periods (N = 63 trials). Solid bars indicate trials '
         'reporting race / ethnicity; hatched bars indicate non-reporting. '
         'Period denominators (n) are shown below each group; percentages '
         'are within-period.'),
        ('figure4_race_by_funder.png',
         'Supplementary Figure S2. Race / ethnicity reporting by funder '
         'type (N = 63 trials). Solid segments indicate trials reporting '
         'race / ethnicity; hatched segments indicate non-reporting. Counts '
         'inside segments, with percent-reporting annotated at right.'),
        ('figure5_indigenous_heatmap.png',
         'Supplementary Figure S3. Indigenous participation and OCAP / '
         'governance indicators across the seven trials that documented '
         'Indigenous participation. Solid circle = Yes; half-filled = '
         'Partial or non-specific; open = No.'),
        ('figure6_sex_gender.png',
         'Supplementary Figure S4. Sex vs. gender reporting in 63 trials. '
         'Sex only = 54 (85.7%); Sex + Gender (not distinguished) = 2 '
         '(3.2%); Sex + Gender (distinguished) = 4 (6.3%); Gender only = '
         '3 (4.8%).'),
        ('figure7_consort_equity.png',
         'Supplementary Figure S5. CONSORT-Equity compliance. Full '
         'compliance was not achieved by any trial (0 / 63). Partial = 17 '
         '(27.0%); None = 46 (73.0%).'),
        ('figure8_bubble_completeness.png',
         'Supplementary Figure S6. PROGRESS-Plus reporting completeness '
         '(0\u201313 fields reported) by publication year and sample size. '
         'Marker shape encodes disorder category; marker area scales with '
         'sample size. The dashed line is the linear trend across all 63 '
         'trials (slope +0.07 / yr; p = 0.28; R\u00B2 = 0.02).'),
    ]
    for fname, caption in fig_list:
        add_page_break(doc)
        add_figure(doc, os.path.join(FIG_DIR, fname),
                   caption, width_cm=15.5)


# ============================================================================
# MAIN
# ============================================================================
def main():
    doc = Document()

    # Page setup: 1-inch margins, A4/letter default
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(1.0)

    # Default paragraph font
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)

    write_title_page(doc)
    write_toc(doc)
    write_s1_checklist(doc)
    write_s2_search(doc)
    write_s3_codebook(doc)
    write_s4_extracted(doc)
    write_s5_irr(doc)
    write_s6_excluded(doc)
    write_figs(doc)

    doc.save(OUT_PATH)
    size_kb = os.path.getsize(OUT_PATH) / 1024
    print(f'Saved: {OUT_PATH} ({size_kb:,.0f} KB)')


if __name__ == '__main__':
    main()
