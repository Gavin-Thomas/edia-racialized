#!/usr/bin/env python3
"""
Generate publication-quality figures and tables for the EDIA scoping review.

Project: EDIA Reporting in Canadian Mental Health Pharmacotherapy RCTs
         (2016-2026)

Run 10 (2026-04-16) -- complete redesign:
  * Pure black-and-white throughout (no decorative color)
  * Times New Roman serif typography
  * Figure 1 (PRISMA-ScR) rendered with Graphviz / dot
  * Figures 2-8 rendered with matplotlib using grayscale fills + hatching
  * All narrative annotations moved out of figures and into APA captions
  * Strict alignment: shared margins, consistent type sizes, axis units

Outputs:
  figures/figure1_prisma_flow.png ... figures/figure8_bubble_completeness.png
  tables_and_figures.docx
"""

import os
import shutil
import subprocess

import numpy as np
import pandas as pd

import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.lines import Line2D
from matplotlib.patches import Rectangle

from scipy import stats
from scipy.stats import linregress

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


# =============================================================================
# CONFIGURATION
# =============================================================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
PROJECT_DIR = os.path.dirname(BASE_DIR)
DATA_CSV = os.path.join(PROJECT_DIR, '06_data_extraction', 'extracted_data.csv')
FIG_DIR = os.path.join(BASE_DIR, 'figures')
DOCX_PATH = os.path.join(BASE_DIR, 'tables_and_figures.docx')
os.makedirs(FIG_DIR, exist_ok=True)


# Pure black-and-white print palette. Differentiation is by fill density and
# hatching, not by hue. All values are grayscale.
PALETTE = {
    'black':     '#000000',
    'fill_dark': '#2E2E2E',   # primary filled bars / emphasis
    'fill_mid':  '#7F7F7F',   # secondary fills
    'fill_lt':   '#BFBFBF',   # tertiary / low-emphasis fills
    'gray_dk':   '#404040',
    'gray_md':   '#808080',
    'gray_lt':   '#CCCCCC',
    'gray_xlt':  '#EAEAEA',   # gridlines / subtle backgrounds
    'text':      '#000000',
    'text_sec':  '#333333',
    'white':     '#FFFFFF',
}


# Global matplotlib settings -- publication-ready Times New Roman, clean spines.
plt.rcParams.update({
    'font.family':         'serif',
    'font.serif':          ['Times New Roman', 'Times', 'DejaVu Serif'],
    'mathtext.fontset':    'stix',
    'font.size':           10.5,
    'axes.titlesize':      12,
    'axes.titleweight':    'bold',
    'axes.titlepad':       12,
    'axes.labelsize':      11,
    'axes.labelcolor':     PALETTE['black'],
    'axes.edgecolor':      PALETTE['black'],
    'axes.linewidth':      0.8,
    'xtick.labelsize':     10,
    'ytick.labelsize':     10,
    'xtick.color':         PALETTE['black'],
    'ytick.color':         PALETTE['black'],
    'xtick.direction':     'out',
    'ytick.direction':     'out',
    'figure.dpi':          300,
    'savefig.dpi':         300,
    'savefig.bbox':        'tight',
    'savefig.pad_inches':  0.20,
    'savefig.facecolor':   'white',
    'axes.spines.top':     False,
    'axes.spines.right':   False,
    'axes.spines.left':    True,
    'axes.spines.bottom':  True,
    'legend.frameon':      True,
    'legend.fancybox':     False,
    'legend.edgecolor':    PALETTE['black'],
    'legend.facecolor':    'white',
    'legend.framealpha':   1.0,
    'legend.fontsize':     9.5,
    'patch.linewidth':     0.6,
    'hatch.linewidth':     0.55,
    'hatch.color':         PALETTE['black'],
    'grid.color':          PALETTE['gray_xlt'],
    'grid.linewidth':      0.5,
})


# =============================================================================
# DATA
# =============================================================================
df_raw = pd.read_csv(DATA_CSV)
df = df_raw[~df_raw['record_number'].isin([4, 12, 45])].copy()
N = len(df)
assert N == 63, f"Expected 63 studies, got {N}"


# =============================================================================
# HELPERS
# =============================================================================
def wilson_ci(k, n, alpha=0.05):
    """Wilson score interval for a binomial proportion (returns % endpoints)."""
    if n == 0:
        return (0.0, 0.0)
    z = stats.norm.ppf(1 - alpha / 2)
    p = k / n
    denom = 1 + z ** 2 / n
    center = (p + z ** 2 / (2 * n)) / denom
    spread = z * np.sqrt(p * (1 - p) / n + z ** 2 / (4 * n ** 2)) / denom
    return (max(0, center - spread) * 100, min(1, center + spread) * 100)


def count_reported(series, include_partial=True):
    if include_partial:
        return series.apply(
            lambda x: str(x).strip().lower() in ('yes', 'partial')).sum()
    return series.apply(lambda x: str(x).strip().lower() == 'yes').sum()


# ----- APA Word-document helpers --------------------------------------------
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, attrs in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{attrs.get("val", "single")}" '
            f'w:sz="{attrs.get("sz", 4)}" w:space="0" '
            f'w:color="{attrs.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def apply_apa_table_borders(table):
    """APA 7 rules: heavy top, thin below header, heavy bottom; no vertical."""
    n_rows = len(table.rows)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            borders = {
                'left':  {"sz": 0, "val": "none", "color": "FFFFFF"},
                'right': {"sz": 0, "val": "none", "color": "FFFFFF"},
            }
            if row_idx == 0:
                borders['top']    = {"sz": 12, "val": "single", "color": "000000"}
                borders['bottom'] = {"sz": 6,  "val": "single", "color": "000000"}
            elif row_idx == n_rows - 1:
                borders['bottom'] = {"sz": 12, "val": "single", "color": "000000"}
                borders['top']    = {"sz": 0, "val": "none", "color": "FFFFFF"}
            else:
                borders['top']    = {"sz": 0, "val": "none", "color": "FFFFFF"}
                borders['bottom'] = {"sz": 0, "val": "none", "color": "FFFFFF"}
            set_cell_border(cell, **borders)


def add_apa_title_block(doc, label, title_text):
    p1 = doc.add_paragraph()
    r = p1.add_run(label); r.bold = True
    r.font.size = Pt(12); r.font.name = 'Times New Roman'
    p1.paragraph_format.space_after = Pt(0)
    p1.paragraph_format.space_before = Pt(12)

    p2 = doc.add_paragraph()
    r = p2.add_run(title_text); r.italic = True
    r.font.size = Pt(12); r.font.name = 'Times New Roman'
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.space_before = Pt(0)


def add_table_note(doc, note_text):
    p = doc.add_paragraph()
    r = p.add_run('Note. '); r.italic = True
    r.font.size = Pt(10); r.font.name = 'Times New Roman'
    r = p.add_run(note_text)
    r.font.size = Pt(10); r.font.name = 'Times New Roman'
    p.paragraph_format.space_before = Pt(4)


def format_table_cell(cell, text, bold=False, size=10, align='left'):
    cell.text = ''
    p = cell.paragraphs[0]
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(str(text))
    r.font.size = Pt(size); r.font.name = 'Times New Roman'
    r.bold = bold
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = Pt(13)


# =============================================================================
# FIGURE 1 -- PRISMA-ScR FLOW (Graphviz / dot)
# =============================================================================
PRISMA_DOT_TEMPLATE = r"""
digraph PRISMA {{
    // ---- Global look ------------------------------------------------------
    graph  [rankdir=TB, splines=ortho, nodesep=0.30, ranksep=0.55,
            bgcolor="white", margin=0.15, fontname="Times New Roman"];
    node   [shape=box, style="filled", fillcolor="white",
            color="black", penwidth=1.0, fontname="Times New Roman",
            fontsize=11, margin="0.18,0.10"];
    edge   [color="black", penwidth=1.0, arrowsize=0.7,
            fontname="Times New Roman", fontsize=10];

    // ---- Row 1: Identification -------------------------------------------
    {{ rank = same;
        stage_id   [label=<<B>IDENTIFICATION</B>>, shape=plaintext,
                    fontsize=12, width=1.4];
        n_id       [label={id_label}, width=3.7];
        n_id_excl  [label={id_excl_label}, width=3.7];
    }}

    // ---- Row 2: Screened (with side exclusion) ---------------------------
    {{ rank = same;
        stage_sc1  [label=<<B>SCREENING</B>>, shape=plaintext,
                    fontsize=12, width=1.4];
        n_screened     [label={screened_label}, width=3.7];
        n_screened_excl[label={screened_excl_label}, width=3.7];
    }}

    // ---- Row 3: Sought for retrieval --------------------------------------
    {{ rank = same;
        stage_sc2  [label="", shape=plaintext, width=1.4];
        n_sought   [label={sought_label}, width=3.7];
        n_notret   [label={notret_label}, width=3.7];
    }}

    // ---- Row 4: Assessed for eligibility ----------------------------------
    {{ rank = same;
        stage_sc3  [label="", shape=plaintext, width=1.4];
        n_assessed     [label={assessed_label}, width=3.7];
        n_assessed_excl[label={assessed_excl_label}, width=3.7];
    }}

    // ---- Row 5: Included --------------------------------------------------
    {{ rank = same;
        stage_in   [label=<<B>INCLUDED</B>>, shape=plaintext,
                    fontsize=12, width=1.4];
        n_included [label={included_label}, width=3.7, penwidth=2.0];
        spacer_in  [label="", shape=plaintext, width=3.7, height=0.01];
    }}

    // ---- Force vertical alignment of the three columns -------------------
    stage_id  -> stage_sc1 -> stage_sc2 -> stage_sc3 -> stage_in
        [style=invis, weight=100];
    n_id_excl -> n_screened_excl -> n_notret -> n_assessed_excl -> spacer_in
        [style=invis, weight=100];

    // ---- Main downward flow (center column) ------------------------------
    n_id        -> n_screened [weight=20];
    n_screened  -> n_sought   [weight=20];
    n_sought    -> n_assessed [weight=20];
    n_assessed  -> n_included [weight=20];

    // ---- Side exclusions (rightward arrows) ------------------------------
    n_id        -> n_id_excl;
    n_screened  -> n_screened_excl;
    n_sought    -> n_notret;
    n_assessed  -> n_assessed_excl;
}}
"""


def _dot_label(title, n_text, sublines=None):
    """Build a Graphviz HTML label with bold count and optional sublines."""
    parts = [title, f'<B>(n = {n_text})</B>']
    if sublines:
        for s in sublines:
            parts.append(f'<FONT POINT-SIZE="10">{s}</FONT>')
    return '<' + '<BR/>'.join(parts) + '>'


def generate_figure1():
    """PRISMA-ScR flow diagram rendered with Graphviz (dot)."""
    dot_path = shutil.which('dot')
    if not dot_path:
        raise RuntimeError("Graphviz `dot` binary not found in PATH.")

    fields = {
        'id_label': _dot_label(
            'Records identified from databases',
            '54,483',
            ['PubMed 9,964 &#183; Europe PMC 15,772',
             'Scopus 27,983 &#183; OpenAlex 764'],
        ),
        'id_excl_label': _dot_label(
            'Records removed before screening',
            '43,579',
            ['Duplicates 14,497 &#183; No abstract 1,557',
             'Out of date 12,473 &#183; Not relevant 12,511',
             'Sys. reviews / meta-analyses 2,541'],
        ),
        'screened_label': _dot_label(
            'Records screened (title / abstract)',
            '10,904',
            ['Dual independent review &#183; &#954; = 0.39'],
        ),
        'screened_excl_label': _dot_label(
            'Records excluded at title / abstract', '10,810'),
        'sought_label': _dot_label('Reports sought for retrieval', '134'),
        'notret_label': _dot_label('Reports not retrieved', '1'),
        'assessed_label': _dot_label('Reports assessed for eligibility', '133'),
        'assessed_excl_label': _dot_label(
            'Reports excluded at full text',
            '70',
            ['Secondary / post-hoc 35',
             'No Canadian site 16',
             'Not an RCT 10',
             'Non-pharma / not mental health 8',
             'Pre-2016 primary 1'],
        ),
        'included_label': _dot_label(
            'Studies included in review',
            '63',
            ['N = 8,837 participants &#183; PROGRESS-Plus extraction'],
        ),
    }

    dot_src = PRISMA_DOT_TEMPLATE.format(**fields)
    dot_file = os.path.join(FIG_DIR, 'figure1_prisma_flow.dot')
    png_file = os.path.join(FIG_DIR, 'figure1_prisma_flow.png')
    with open(dot_file, 'w', encoding='utf-8') as f:
        f.write(dot_src)

    res = subprocess.run(
        [dot_path, '-Tpng', '-Gdpi=300', dot_file, '-o', png_file],
        capture_output=True, text=True,
    )
    if res.returncode != 0:
        raise RuntimeError(f"dot failed: {res.stderr}")

    print("  Figure 1: PRISMA flow diagram (Graphviz) saved.")


# =============================================================================
# TABLE 1 DATA
# =============================================================================
def generate_table1_data():
    rows = []
    year_bins = [(2013, 2017), (2018, 2020), (2021, 2023), (2024, 2026)]

    rows.append(('Year of publication', ''))
    for lo, hi in year_bins:
        n = ((df['year'] >= lo) & (df['year'] <= hi)).sum()
        rows.append((f'  {lo}\u2013{hi}', f'{n} ({n/N*100:.1f})'))

    rows.append(('Disorder category', ''))
    disorder_order = ['Depression', 'Dementia', 'Bipolar-Mood', 'Substance Use',
                      'Psychotic', 'ADHD', 'Anxiety-Trauma-OCD',
                      'Eating Disorder', 'Other']
    for d in disorder_order:
        n = (df['disorder_category'] == d).sum()
        if n > 0:
            rows.append((f'  {d}', f'{n} ({n/N*100:.1f})'))

    rows.append(('Trial design', ''))
    for d in ['Parallel', 'Crossover', 'Other']:
        if d == 'Other':
            n = (~df['trial_design'].isin(['Parallel', 'Crossover'])).sum()
        else:
            n = (df['trial_design'] == d).sum()
        rows.append((f'  {d}', f'{n} ({n/N*100:.1f})'))

    rows.append(('Setting', ''))
    n_multi  = (df['multisite']     == 'Yes').sum()
    n_single = (df['multisite']     == 'No').sum()
    n_intl   = (df['international'] == 'Yes').sum()
    n_can    = (df['international'] == 'No').sum()
    rows.append(('  Multisite',       f'{n_multi} ({n_multi/N*100:.1f})'))
    rows.append(('  Single-site',     f'{n_single} ({n_single/N*100:.1f})'))
    rows.append(('  International',   f'{n_intl} ({n_intl/N*100:.1f})'))
    rows.append(('  All-Canadian',    f'{n_can} ({n_can/N*100:.1f})'))

    rows.append(('Sample size', ''))
    med = df['sample_size'].median()
    q1  = df['sample_size'].quantile(0.25)
    q3  = df['sample_size'].quantile(0.75)
    mn  = df['sample_size'].min()
    mx  = df['sample_size'].max()
    rows.append(('  Median [IQR]', f'{med:.0f} [{q1:.0f}\u2013{q3:.0f}]'))
    rows.append(('  Range', f'{mn:.0f}\u2013{mx:.0f}'))

    rows.append(('Funder type', ''))
    for f in ['Industry', 'CIHR', 'NIH', 'Other government', 'Foundation',
              'Not reported']:
        n = (df['funder'] == f).sum()
        if n > 0:
            rows.append((f'  {f}', f'{n} ({n/N*100:.1f})'))
    return rows


# =============================================================================
# FIGURE 2 -- PROGRESS-Plus reporting rates (B&W horizontal bars)
# =============================================================================
def generate_figure2():
    variables = [
        ('Age',                    count_reported(df['age_reported'])),
        ('Place of residence',     count_reported(df['place_reported'])),
        ('Sex',                    count_reported(df['sex_reported'])),
        ('Race / Ethnicity',
            (df['race_reported'].str.strip().str.lower() == 'yes').sum()),
        ('Disability',             count_reported(df['disability_reported'])),
        ('Education',              count_reported(df['education_reported'])),
        ('SES / Income',           count_reported(df['ses_reported'])),
        ('Occupation',             count_reported(df['occupation_reported'])),
        ('Gender',                 count_reported(df['gender_reported'])),
        ('SOGI',                   count_reported(df['sogi_reported'])),
        ('Social capital',         count_reported(df['social_capital_reported'])),
        ('Religion',               count_reported(df['religion_reported'])),
        ('Intersectional analysis',count_reported(df['intersectional_analysis'])),
    ]

    variables_sorted = sorted(variables, key=lambda x: x[1])
    labels = [v[0] for v in variables_sorted]
    values = [v[1] for v in variables_sorted]
    pcts   = [v / N * 100 for v in values]

    fig, ax = plt.subplots(figsize=(7.5, 5.4))

    # Bars: solid black fill, no hatching -- the variable name carries meaning.
    ax.barh(labels, pcts,
            color=PALETTE['fill_dark'],
            edgecolor=PALETTE['black'], linewidth=0.6,
            height=0.66, zorder=3)

    # Right-side data labels: count and percent.
    for i, (n_val, pct) in enumerate(zip(values, pcts)):
        ax.text(pct + 1.5, i, f'{n_val} / 63   ({pct:.1f}%)',
                va='center', ha='left', fontsize=9.5,
                color=PALETTE['black'])

    # Subtle reference rule at 50%.
    ax.axvline(50, color=PALETTE['gray_md'], linewidth=0.7,
               linestyle=(0, (4, 3)), zorder=1)
    ax.text(50, len(labels) - 0.3, '  50%', fontsize=9,
            color=PALETTE['gray_dk'], va='center', ha='left')

    ax.set_xlim(0, 118)
    ax.set_xticks([0, 25, 50, 75, 100])
    ax.set_xlabel('Trials reporting (%)')
    ax.set_title('PROGRESS-Plus Variable Reporting Rates (N = 63 Trials)',
                 loc='left')
    ax.set_axisbelow(True)
    ax.xaxis.grid(True, alpha=0.5)
    ax.tick_params(axis='y', length=0)

    plt.tight_layout()
    plt.savefig(os.path.join(FIG_DIR, 'figure2_progress_plus_bars.png'))
    plt.close()
    print("  Figure 2: PROGRESS-Plus bars saved.")


# =============================================================================
# FIGURE 3 -- Race/ethnicity reporting over time (B&W grouped bars)
# =============================================================================
def generate_figure3():
    def year_bin(y):
        if y <= 2017: return '2013\u20132017'
        if y <= 2019: return '2018\u20132019'
        if y <= 2021: return '2020\u20132021'
        if y <= 2023: return '2022\u20132023'
        return '2024\u20132026'

    dfp = df.copy()
    dfp['year_bin'] = dfp['year'].apply(year_bin)
    dfp['race_yes'] = dfp['race_reported'].str.strip().str.lower() == 'yes'
    bin_order = ['2013\u20132017', '2018\u20132019', '2020\u20132021',
                 '2022\u20132023', '2024\u20132026']
    grouped = dfp.groupby('year_bin').agg(
        total=('race_yes', 'count'),
        reported=('race_yes', 'sum'),
    ).reindex(bin_order)
    grouped['not_reported'] = grouped['total'] - grouped['reported']
    grouped['pct_reported'] = grouped['reported'] / grouped['total'] * 100

    fig, ax = plt.subplots(figsize=(8.0, 4.8))
    x = np.arange(len(bin_order))
    width = 0.36

    # Reported -- solid black; Not reported -- white with diagonal hatching.
    ax.bar(x - width / 2, grouped['pct_reported'], width,
           color=PALETTE['fill_dark'], edgecolor=PALETTE['black'],
           linewidth=0.7, label='Reported', zorder=3)
    ax.bar(x + width / 2, 100 - grouped['pct_reported'], width,
           color='white', edgecolor=PALETTE['black'], linewidth=0.7,
           hatch='////', label='Not reported', zorder=3)

    # Per-bar labels above bars.
    for i in range(len(bin_order)):
        n_rep = int(grouped['reported'].iloc[i])
        n_not = int(grouped['not_reported'].iloc[i])
        pct_r = grouped['pct_reported'].iloc[i]
        ax.text(x[i] - width / 2, pct_r + 1.5,
                f'{n_rep}\n({pct_r:.0f}%)',
                ha='center', va='bottom', fontsize=9, color=PALETTE['black'])
        ax.text(x[i] + width / 2, 100 - pct_r + 1.5,
                f'{n_not}\n({100 - pct_r:.0f}%)',
                ha='center', va='bottom', fontsize=9,
                color=PALETTE['text_sec'])

    # Period denominators below the x-axis.
    for i in range(len(bin_order)):
        ax.annotate(f'n = {int(grouped["total"].iloc[i])}',
                    xy=(i, 0), xycoords=('data', 'axes fraction'),
                    xytext=(0, -28), textcoords='offset points',
                    ha='center', fontsize=9, color=PALETTE['text_sec'])

    ax.set_xticks(x)
    ax.set_xticklabels(bin_order)
    ax.set_yticks([0, 25, 50, 75, 100])
    ax.set_ylim(0, 110)
    ax.set_xlabel('Publication period', labelpad=20)
    ax.set_ylabel('Trials in period (%)')
    ax.set_title('Race / Ethnicity Reporting Over Time (N = 63 Trials)',
                 loc='left')
    ax.set_axisbelow(True)
    ax.yaxis.grid(True, alpha=0.5)
    ax.legend(loc='upper left', bbox_to_anchor=(0.0, 1.0))

    plt.tight_layout()
    plt.savefig(os.path.join(FIG_DIR, 'figure3_race_over_time.png'))
    plt.close()
    print("  Figure 3: Race / ethnicity over time saved.")


# =============================================================================
# TABLE 2 DATA
# =============================================================================
def generate_table2_data():
    rows = []

    def add_row(label, n_reported, total=N, indent=False):
        pct = n_reported / total * 100
        ci_lo, ci_hi = wilson_ci(n_reported, total)
        prefix = '  ' if indent else ''
        rows.append((f'{prefix}{label}',
                     f'{n_reported}/{total}',
                     f'{pct:.1f}',
                     f'{ci_lo:.1f}\u2013{ci_hi:.1f}'))

    rows.append(('PROGRESS-Plus Core', '', '', ''))
    add_row('Place of residence (P)', count_reported(df['place_reported']))
    n_race = (df['race_reported'].str.strip().str.lower() == 'yes').sum()
    add_row('Race/ethnicity (R)', n_race)

    race_df = df[df['race_reported'].str.strip().str.lower() == 'yes']
    n_race_total = len(race_df)
    rows.append(('  Framework (among reporters)', '', '', ''))
    for fw in ['US-derived', 'Canadian Census', 'Self-identified',
               'Other', 'Not stated']:
        n_fw = (race_df['race_framework'].str.strip() == fw).sum()
        if n_fw > 0:
            rows.append((f'    {fw}', f'{n_fw}/{n_race_total}',
                         f'{n_fw/n_race_total*100:.1f}', '\u2014'))

    rows.append(('  Granularity (among reporters)', '', '', ''))
    for g in ['Granular', 'Moderate', 'Coarse', 'Binary']:
        n_g = (race_df['race_granularity'].str.strip() == g).sum()
        if n_g > 0:
            rows.append((f'    {g}', f'{n_g}/{n_race_total}',
                         f'{n_g/n_race_total*100:.1f}', '\u2014'))

    add_row('Subgroup analysis conducted',
            (df['race_subgroup_analysis'].str.strip().str.lower() == 'yes').sum(),
            indent=True)
    n_cov = (df['race_as_variable'].str.strip().isin(
        ['Covariate', 'Subgroup'])).sum()
    add_row('Used as covariate/subgroup', n_cov, indent=True)
    add_row('Acknowledged as social construct',
            (df['race_social_construct'].str.strip().str.lower() == 'yes').sum(),
            indent=True)

    add_row('Occupation (O)', count_reported(df['occupation_reported']))
    add_row('Sex (G)', count_reported(df['sex_reported']))
    add_row('Gender (G)', count_reported(df['gender_reported']))
    add_row('Sex/gender distinguished',
            (df['sex_gender_distinguished'].str.strip().str.lower() == 'yes').sum(),
            indent=True)
    add_row('Religion (R)', count_reported(df['religion_reported']))
    add_row('Education (E)', count_reported(df['education_reported']))
    add_row('SES / income (S)', count_reported(df['ses_reported']))
    add_row('Social capital (S)', count_reported(df['social_capital_reported']))

    rows.append(('PROGRESS "Plus" Variables', '', '', ''))
    add_row('Age',                     count_reported(df['age_reported']))
    add_row('Disability',              count_reported(df['disability_reported']))
    add_row('SOGI',                    count_reported(df['sogi_reported']))
    add_row('Intersectional analysis', count_reported(df['intersectional_analysis']))
    return rows


# =============================================================================
# FIGURE 4 -- Race/ethnicity reporting by funder type (stacked horizontal)
# =============================================================================
def generate_figure4():
    dfp = df.copy()
    dfp['race_yes'] = dfp['race_reported'].str.strip().str.lower() == 'yes'

    funder_order = ['NIH', 'CIHR', 'Other government', 'Industry',
                    'Foundation', 'Not reported']
    rows = []
    for f in funder_order:
        sub = dfp[dfp['funder'] == f]
        if len(sub) == 0:
            continue
        n_total = len(sub)
        n_yes = int(sub['race_yes'].sum())
        rows.append(dict(funder=f, reported=n_yes,
                         not_reported=n_total - n_yes,
                         total=n_total, pct=n_yes / n_total * 100))
    fd = pd.DataFrame(rows)

    fig, ax = plt.subplots(figsize=(8.0, 4.6))
    y = np.arange(len(fd))

    ax.barh(y, fd['reported'], height=0.55,
            color=PALETTE['fill_dark'], edgecolor=PALETTE['black'],
            linewidth=0.7, label='Race reported', zorder=3)
    ax.barh(y, fd['not_reported'], left=fd['reported'], height=0.55,
            color='white', edgecolor=PALETTE['black'], linewidth=0.7,
            hatch='////', label='Race not reported', zorder=3)

    bbox_lbl = dict(facecolor='white', edgecolor='none',
                    pad=1.5, boxstyle='square,pad=0.15')
    for i, r in fd.iterrows():
        if r['reported'] >= 2:
            ax.text(r['reported'] / 2, i, f"{int(r['reported'])}",
                    ha='center', va='center', fontsize=10,
                    color='white', fontweight='bold')
        if r['not_reported'] >= 2:
            # White bbox keeps the count legible against diagonal hatching.
            ax.text(r['reported'] + r['not_reported'] / 2, i,
                    f"{int(r['not_reported'])}",
                    ha='center', va='center', fontsize=10,
                    color=PALETTE['black'], bbox=bbox_lbl)
        ax.text(r['total'] + 0.5, i, f"{r['pct']:.0f}% reported",
                va='center', ha='left', fontsize=9.5,
                color=PALETTE['text_sec'])

    ax.set_yticks(y)
    ax.set_yticklabels(fd['funder'])
    ax.invert_yaxis()
    ax.set_xlim(0, fd['total'].max() + 7)
    ax.set_xlabel('Number of trials')
    ax.set_title('Race / Ethnicity Reporting by Funder Type (N = 63 Trials)',
                 loc='left', pad=14)
    ax.set_axisbelow(True)
    ax.xaxis.grid(True, alpha=0.5)
    ax.tick_params(axis='y', length=0)
    # Legend below the plot to avoid colliding with the title.
    ax.legend(loc='upper center', bbox_to_anchor=(0.5, -0.18), ncol=2)

    plt.tight_layout()
    plt.savefig(os.path.join(FIG_DIR, 'figure4_race_by_funder.png'))
    plt.close()
    print("  Figure 4: Race by funder saved.")


# =============================================================================
# TABLE 3 DATA
# =============================================================================
def generate_table3_data():
    funder_order = ['Industry', 'CIHR', 'NIH', 'Other government',
                    'Foundation', 'Not reported']
    rows = []
    contingency = []
    for f in funder_order:
        sub = df[df['funder'] == f]
        n_total = len(sub)
        if n_total == 0:
            continue
        n_yes     = (sub['race_reported'].str.strip().str.lower() == 'yes').sum()
        n_partial = (sub['race_reported'].str.strip().str.lower() == 'partial').sum()
        n_no      = (sub['race_reported'].str.strip().str.lower() == 'no').sum()
        rows.append((f, str(n_yes), str(n_partial), str(n_no), str(n_total),
                     f'{n_yes/n_total*100:.1f}'))
        contingency.append([n_yes, n_no])
    try:
        chi2, p_val, dof, _ = stats.chi2_contingency(np.array(contingency))
        p_text = f'Chi-square = {chi2:.2f}, df = {dof}, p = {p_val:.3f}'
    except Exception:
        p_text = 'Statistical test not computed.'
    return rows, p_text


# =============================================================================
# FIGURE 5 -- Indigenous participation and governance (B&W dot matrix)
# =============================================================================
def generate_figure5():
    indig = df[df['indigenous_participation'] == 'Yes'].copy()

    indicators = [
        ('Participation\ndocumented',   'indigenous_participation'),
        ('Groups\nnamed',               'indigenous_groups_named'),
        ('Indigenous-\nspecific trial', 'indigenous_specific_trial'),
        ('Separate\nreporting',         'indigenous_separate_reporting'),
        ('Community\npartnership',      'indigenous_partnership'),
        ('Data ownership\ndescribed',   'indigenous_data_ownership'),
        ('Governance\nbody',            'indigenous_governance_body'),
    ]

    # Disambiguate duplicate author/year combos as Shiwach 2025a / 2025b.
    ay_counts = indig.groupby(['first_author', 'year']).cumcount()
    ay_totals = indig.groupby(['first_author', 'year'])['first_author'] \
                     .transform('count')

    trial_labels, matrix = [], []
    for idx_i, (_, row) in enumerate(indig.iterrows()):
        base = f"{row['first_author']} ({row['year']})"
        if ay_totals.iloc[idx_i] > 1:
            base = f"{row['first_author']} ({row['year']}{chr(ord('a') + ay_counts.iloc[idx_i])})"
        trial_labels.append(base)

        vals = []
        for _, col in indicators:
            v = str(row[col]).strip().lower()
            if v == 'yes':
                vals.append(1.0)
            elif v in ('no', 'not reported', '0', 'nan', ''):
                vals.append(0.0)
            elif v == 'partial' or 'not specified' in v or 'not further' in v:
                vals.append(0.5)
            else:
                vals.append(1.0 if v else 0.0)
        matrix.append(vals)

    # Override "groups named" coding from raw text.
    for i, (_, row) in enumerate(indig.iterrows()):
        gn = str(row['indigenous_groups_named']).strip().lower()
        if not gn or gn in ('nan', 'no', 'not specified', 'not reported'):
            matrix[i][1] = 0.0
        elif any(k in gn for k in ('first nations', 'metis', 'inuit')):
            matrix[i][1] = 1.0
        else:
            matrix[i][1] = 0.5

    matrix = np.array(matrix)
    col_labels = [i[0] for i in indicators]
    n_rows, n_cols = matrix.shape

    fig, ax = plt.subplots(figsize=(9.0, 5.0))

    # Light gray gridlines.
    for i in range(n_rows + 1):
        ax.axhline(y=i - 0.5, color=PALETTE['gray_xlt'],
                   linewidth=0.5, zorder=1)
    for j in range(n_cols + 1):
        ax.axvline(x=j - 0.5, color=PALETTE['gray_xlt'],
                   linewidth=0.5, zorder=1)

    # Vertical separator between Documented (cols 0-1) and Governance (cols 2-6).
    ax.axvline(x=1.5, color=PALETTE['black'], linewidth=1.0, zorder=2)

    # Draw markers: filled black = Yes, half-fill = Partial, open = No.
    ms = 18
    for i in range(n_rows):
        for j in range(n_cols):
            v = matrix[i, j]
            if v >= 0.75:        # Yes -- solid black
                ax.plot(j, i, 'o', markersize=ms,
                        markerfacecolor=PALETTE['black'],
                        markeredgecolor=PALETTE['black'],
                        markeredgewidth=0.8, zorder=5)
            elif v >= 0.25:      # Partial -- gray with stroke
                ax.plot(j, i, 'o', markersize=ms,
                        markerfacecolor=PALETTE['gray_md'],
                        markeredgecolor=PALETTE['black'],
                        markeredgewidth=0.8, zorder=5)
            else:                # No -- open ring
                ax.plot(j, i, 'o', markersize=ms,
                        markerfacecolor='white',
                        markeredgecolor=PALETTE['black'],
                        markeredgewidth=0.8, zorder=5)

    # Column group banners above the headers.
    ax.text(0.5, -1.55, 'DOCUMENTED', ha='center', va='center',
            fontsize=9.5, fontweight='bold', color=PALETTE['black'])
    ax.text(4.0, -1.55, 'GOVERNANCE & OCAP INDICATORS',
            ha='center', va='center', fontsize=9.5, fontweight='bold',
            color=PALETTE['black'])

    ax.set_xlim(-0.7, n_cols - 0.3)
    ax.set_ylim(n_rows - 0.5, -2.1)
    ax.set_xticks(range(n_cols))
    ax.set_xticklabels(col_labels, fontsize=9.5, ha='center')
    ax.xaxis.set_ticks_position('top')
    ax.xaxis.set_label_position('top')
    ax.set_yticks(range(n_rows))
    ax.set_yticklabels(trial_labels, fontsize=10)
    for spine in ax.spines.values():
        spine.set_visible(False)
    ax.tick_params(top=False, bottom=False, left=False, right=False)

    ax.set_title('Indigenous Participation and Governance Indicators '
                 '(7 Trials)', loc='center', pad=42)

    legend_elements = [
        Line2D([0], [0], marker='o', color='w', markersize=10,
               markerfacecolor=PALETTE['black'],
               markeredgecolor=PALETTE['black'],
               markeredgewidth=0.8, label='Yes'),
        Line2D([0], [0], marker='o', color='w', markersize=10,
               markerfacecolor=PALETTE['gray_md'],
               markeredgecolor=PALETTE['black'],
               markeredgewidth=0.8, label='Partial / non-specific'),
        Line2D([0], [0], marker='o', color='w', markersize=10,
               markerfacecolor='white',
               markeredgecolor=PALETTE['black'],
               markeredgewidth=0.8, label='No'),
    ]
    ax.legend(handles=legend_elements, loc='lower center',
              bbox_to_anchor=(0.5, -0.20), ncol=3)

    plt.tight_layout()
    plt.savefig(os.path.join(FIG_DIR, 'figure5_indigenous_heatmap.png'))
    plt.close()
    print("  Figure 5: Indigenous dot matrix saved.")


# =============================================================================
# TABLE 4 DATA
# =============================================================================
def generate_table4_data():
    indig = df[df['indigenous_participation'] == 'Yes'].copy()
    ay_counts = indig.groupby(['first_author', 'year']).cumcount()
    ay_totals = indig.groupby(['first_author', 'year'])['first_author'] \
                     .transform('count')
    rows = []
    for idx_i, (_, r) in enumerate(indig.iterrows()):
        label = f"{r['first_author']} ({r['year']})"
        if ay_totals.iloc[idx_i] > 1:
            label = (f"{r['first_author']} "
                     f"({r['year']}{chr(ord('a') + ay_counts.iloc[idx_i])})")
        notes = str(r['indigenous_notes'])
        if len(notes) > 150:
            notes = notes[:150] + '\u2026'
        rows.append((
            str(int(r['record_number'])),
            label,
            str(r['disorder_category']),
            str(int(r['sample_size'])),
            str(r['indigenous_groups_named']),
            str(r['indigenous_partnership']),
            str(int(r['indigenous_data_sovereignty'])),
            notes,
        ))
    return rows


# =============================================================================
# FIGURE 6 -- Sex vs gender reporting (B&W stacked horizontal bar)
# =============================================================================
def generate_figure6():
    sex_yes       = df['sex_reported'].str.strip().str.lower().isin(['yes', 'partial'])
    gender_yes    = df['gender_reported'].str.strip().str.lower().isin(['yes', 'partial'])
    distinguished = df['sex_gender_distinguished'].str.strip().str.lower() == 'yes'

    sex_only            = int((sex_yes & ~gender_yes).sum())
    sex_gender_not_dist = int((sex_yes & gender_yes & ~distinguished).sum())
    sex_gender_dist     = int((sex_yes & gender_yes & distinguished).sum())
    gender_only         = int((~sex_yes & gender_yes).sum())
    neither             = int((~sex_yes & ~gender_yes).sum())

    # Each segment differentiated by hatching pattern (B&W).
    segments = [
        ('Sex only',                          sex_only,            PALETTE['fill_dark'], None),
        ('Sex + Gender (not distinguished)',  sex_gender_not_dist, PALETTE['fill_mid'],  None),
        ('Sex + Gender (distinguished)',      sex_gender_dist,     'white',              '////'),
        ('Gender only',                       gender_only,         'white',              'xxxx'),
    ]
    if neither > 0:
        segments.append(('Neither', neither, PALETTE['fill_lt'], None))

    fig, ax = plt.subplots(figsize=(8.5, 2.2))
    left = 0.0
    bar_h = 0.70
    for label, n_val, color, hatch in segments:
        if n_val == 0:
            continue
        pct = n_val / N * 100
        ax.barh(0, pct, left=left, height=bar_h,
                color=color, hatch=hatch,
                edgecolor=PALETTE['black'], linewidth=0.8, zorder=3)
        mid = left + pct / 2
        if pct >= 10:
            text_color = 'white' if color == PALETTE['fill_dark'] else PALETTE['black']
            ax.text(mid, 0, f'{n_val}  ({pct:.1f}%)',
                    ha='center', va='center', fontsize=10,
                    fontweight='bold', color=text_color)
        # Small segments: count is shown in the legend below; no in-bar callouts.
        left += pct

    legend_patches = [
        mpatches.Patch(facecolor=c, hatch=h, edgecolor=PALETTE['black'],
                       linewidth=0.7,
                       label=f'{lbl}  ({n}, {n / N * 100:.1f}%)')
        for lbl, n, c, h in segments if n > 0
    ]
    ax.legend(handles=legend_patches, loc='upper center',
              bbox_to_anchor=(0.5, -0.95), ncol=2, frameon=True,
              handlelength=2.0, handleheight=1.0, columnspacing=2.0)

    ax.set_xlim(0, 100)
    ax.set_ylim(-0.5, 0.5)
    ax.set_yticks([])
    ax.set_xticks([0, 25, 50, 75, 100])
    ax.set_xlabel('Trials (%)', labelpad=6)
    ax.set_title('Sex vs. Gender Reporting (N = 63 Trials)', loc='left', pad=10)
    for s in ('left', 'top', 'right'):
        ax.spines[s].set_visible(False)

    plt.subplots_adjust(top=0.80, bottom=0.55, left=0.05, right=0.97)
    plt.savefig(os.path.join(FIG_DIR, 'figure6_sex_gender.png'))
    plt.close()
    print("  Figure 6: Sex vs gender saved.")


# =============================================================================
# FIGURE 7 -- CONSORT-Equity compliance (B&W stacked horizontal bar)
# =============================================================================
def generate_figure7():
    full    = int((df['consort_equity_compliant'].str.strip().str.lower() == 'yes').sum())
    partial = int((df['consort_equity_compliant'].str.strip().str.lower() == 'partial').sum())
    none_v  = int((df['consort_equity_compliant'].str.strip().str.lower() == 'no').sum())

    segments = [
        ('Full',    full,    'white',                '////'),
        ('Partial', partial, PALETTE['fill_mid'],    None),
        ('None',    none_v,  PALETTE['fill_dark'],   None),
    ]

    fig, ax = plt.subplots(figsize=(8.5, 1.9))
    left = 0.0
    bar_h = 0.70
    for label, n_val, color, hatch in segments:
        pct = n_val / N * 100
        if n_val == 0:
            continue
        ax.barh(0, pct, left=left, height=bar_h,
                color=color, hatch=hatch,
                edgecolor=PALETTE['black'], linewidth=0.8, zorder=3)
        mid = left + pct / 2
        text_color = 'white' if color == PALETTE['fill_dark'] else PALETTE['black']
        if pct >= 15:
            ax.text(mid, 0, f'{label}\nn = {n_val}  ({pct:.1f}%)',
                    ha='center', va='center', fontsize=10.5,
                    fontweight='bold', color=text_color)
        else:
            ax.text(mid, bar_h / 2 + 0.10, f'{label}\nn = {n_val} ({pct:.1f}%)',
                    ha='center', va='bottom', fontsize=9.5,
                    color=PALETTE['black'])
        left += pct

    ax.set_xlim(0, 100)
    ax.set_ylim(-0.5, 0.5)
    ax.set_yticks([])
    ax.set_xticks([0, 25, 50, 75, 100])
    ax.set_xlabel('Trials (%)', labelpad=6)
    ax.set_title('CONSORT-Equity Compliance (N = 63 Trials)', loc='left', pad=10)
    for s in ('left', 'top', 'right'):
        ax.spines[s].set_visible(False)

    plt.subplots_adjust(top=0.78, bottom=0.30, left=0.05, right=0.97)
    plt.savefig(os.path.join(FIG_DIR, 'figure7_consort_equity.png'))
    plt.close()
    print("  Figure 7: CONSORT-Equity bar saved.")


# =============================================================================
# FIGURE 8 -- PROGRESS-Plus completeness by year and sample size
# =============================================================================
def generate_figure8():
    progress_cols = [
        'place_reported', 'race_reported', 'occupation_reported', 'sex_reported',
        'gender_reported', 'religion_reported', 'education_reported',
        'ses_reported', 'social_capital_reported', 'age_reported',
        'disability_reported', 'sogi_reported', 'intersectional_analysis',
    ]

    def compute_score(row):
        score = 0
        for c in progress_cols:
            v = str(row[c]).strip().lower()
            if v == 'yes':
                score += 1
            elif c != 'race_reported' and v == 'partial':
                score += 1
        return score

    dfp = df.copy()
    dfp['completeness'] = dfp.apply(compute_score, axis=1)

    disorder_abbrev = {
        'Anxiety-Trauma-OCD': 'Anxiety / OCD',
        'Bipolar-Mood':       'Bipolar / Mood',
        'Substance Use':      'Substance use',
        'Psychotic':          'Psychotic',
        'Eating Disorder':    'Eating disorder',
        'Depression':         'Depression',
        'Dementia':           'Dementia',
        'ADHD':               'ADHD',
        'Other':              'Other',
    }

    # Distinct B&W marker shapes per disorder category (no color encoding).
    disorder_cats = sorted(dfp['disorder_category'].unique(), key=str)
    marker_cycle = ['o', 's', '^', 'D', 'v', 'P', 'X', '<', '>']
    marker_map = {d: marker_cycle[i % len(marker_cycle)]
                  for i, d in enumerate(disorder_cats)}

    fig, ax = plt.subplots(figsize=(13.0, 6.0))
    np.random.seed(42)
    jitter = np.random.uniform(-0.22, 0.22, size=len(dfp))

    for disorder in disorder_cats:
        mask = dfp['disorder_category'] == disorder
        sub = dfp[mask]
        idx = sub.index
        ax.scatter(sub['year'] + jitter[dfp.index.get_indexer(idx)],
                   sub['completeness'],
                   s=np.clip(sub['sample_size'] * 1.0, 20, 600),
                   marker=marker_map[disorder],
                   facecolor=PALETTE['gray_lt'],
                   edgecolor=PALETTE['black'],
                   linewidth=0.7, alpha=0.85,
                   label=disorder_abbrev.get(disorder, disorder), zorder=3)

    slope, intercept, r_val, p_val, _ = linregress(
        dfp['year'], dfp['completeness'])
    x_line = np.linspace(dfp['year'].min(), dfp['year'].max(), 100)
    ax.plot(x_line, slope * x_line + intercept,
            linestyle='--', color=PALETTE['black'], linewidth=1.4,
            zorder=2,
            label='Linear trend')

    # Annotate the trend statistics in the upper-left of the axes (inside the
    # plotting area) so the legend stays compact and nothing gets clipped.
    trend_stats = (f'Slope = +{slope:.2f}/yr\n'
                   f'p = {p_val:.2f}\n'
                   f'R\u00B2 = {r_val**2:.2f}')
    ax.text(0.015, 0.975, trend_stats, transform=ax.transAxes,
            ha='left', va='top', fontsize=9,
            bbox=dict(boxstyle='round,pad=0.35',
                      facecolor='white',
                      edgecolor=PALETTE['black'],
                      linewidth=0.6))

    ax.set_xlabel('Publication year')
    ax.set_ylabel('PROGRESS-Plus completeness (0\u201313)')
    ax.set_title('PROGRESS-Plus Reporting Completeness by Year and Sample Size '
                 '(N = 63)', loc='left')
    ymax = int(dfp['completeness'].max()) + 2
    ax.set_ylim(0, min(ymax, 13))
    ax.set_yticks(range(0, min(ymax, 13) + 1))
    ax.set_axisbelow(True)
    ax.grid(True, alpha=0.5)

    handles, labels = ax.get_legend_handles_labels()
    cat_legend = ax.legend(handles=handles, labels=labels,
                           loc='upper left', bbox_to_anchor=(1.02, 1.0),
                           title='Disorder category', title_fontsize=10,
                           fontsize=9, borderaxespad=0.0)
    cat_legend.get_title().set_fontweight('bold')
    ax.add_artist(cat_legend)

    size_handles = [
        plt.scatter([], [], s=20,  marker='o',
                    facecolor=PALETTE['gray_lt'],
                    edgecolor=PALETTE['black'], linewidth=0.6, label='20'),
        plt.scatter([], [], s=100, marker='o',
                    facecolor=PALETTE['gray_lt'],
                    edgecolor=PALETTE['black'], linewidth=0.6, label='100'),
        plt.scatter([], [], s=500, marker='o',
                    facecolor=PALETTE['gray_lt'],
                    edgecolor=PALETTE['black'], linewidth=0.6, label='500'),
    ]
    size_legend = ax.legend(handles=size_handles, loc='lower left',
                            bbox_to_anchor=(1.02, 0.0),
                            title='Sample size (N)', title_fontsize=10,
                            fontsize=9, labelspacing=1.4, borderpad=0.9,
                            borderaxespad=0.0)
    size_legend.get_title().set_fontweight('bold')

    plt.tight_layout(rect=[0, 0, 0.74, 1])
    plt.savefig(os.path.join(FIG_DIR, 'figure8_bubble_completeness.png'))
    plt.close()
    print("  Figure 8: Bubble chart saved.")


# =============================================================================
# WORD DOCUMENT (tables_and_figures.docx)
# =============================================================================
def build_docx():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Title page
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Tables and Figures'); r.bold = True
    r.font.size = Pt(14); r.font.name = 'Times New Roman'

    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('EDIA Reporting in Canadian Mental Health Pharmacotherapy '
                  'RCTs (2016\u20132026):\n'
                  'A Scoping Review Using the PROGRESS-Plus Framework')
    r.font.size = Pt(12); r.font.name = 'Times New Roman'; r.italic = True
    doc.add_page_break()

    # ---- Figure 1 ----
    add_apa_title_block(doc, 'Figure 1',
                        'PRISMA-ScR Flow Diagram for Study Selection')
    doc.add_picture(os.path.join(FIG_DIR, 'figure1_prisma_flow.png'),
                    width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table_note(doc,
        'Flow diagram follows the PRISMA-ScR (2020) reporting guidance. '
        'N = 63 studies were retained after excluding three records '
        '(#4, #12, #45) on post-extraction review.')
    doc.add_page_break()

    # ---- Table 1 ----
    add_apa_title_block(doc, 'Table 1',
                        'Characteristics of Included Trials (N = 63)')
    t1 = generate_table1_data()
    table1 = doc.add_table(rows=len(t1) + 1, cols=2)
    table1.alignment = WD_TABLE_ALIGNMENT.CENTER
    format_table_cell(table1.rows[0].cells[0], 'Characteristic', bold=True)
    format_table_cell(table1.rows[0].cells[1], 'n (%) or Median [IQR]',
                      bold=True, align='center')
    for i, (label, value) in enumerate(t1):
        is_section = value == ''
        format_table_cell(table1.rows[i + 1].cells[0], label, bold=is_section)
        format_table_cell(table1.rows[i + 1].cells[1], value, align='center')
    apply_apa_table_borders(table1)
    add_table_note(doc,
        'Values are n (%) unless otherwise specified. IQR = interquartile '
        'range. CIHR = Canadian Institutes of Health Research. NIH = '
        'National Institutes of Health. Record #8 (Malla 2013) was '
        'published before 2016 but met all other inclusion criteria.')
    doc.add_page_break()

    # ---- Figure 2 ----
    add_apa_title_block(doc, 'Figure 2',
        'PROGRESS-Plus Variable Reporting Rates (N = 63 Trials)')
    doc.add_picture(os.path.join(FIG_DIR, 'figure2_progress_plus_bars.png'),
                    width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table_note(doc,
        'Bars show the proportion of trials reporting each PROGRESS-Plus '
        'variable. Race/ethnicity counts strict "Yes" responses only; all '
        'other variables count "Yes" + "Partial". The dashed reference line '
        'marks the 50% threshold.')
    doc.add_page_break()

    # ---- Figure 3 ----
    add_apa_title_block(doc, 'Figure 3',
                        'Race / Ethnicity Reporting Over Time')
    doc.add_picture(os.path.join(FIG_DIR, 'figure3_race_over_time.png'),
                    width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table_note(doc,
        'Years are binned to accommodate small per-year sample sizes. '
        'Bar labels show count and percent of trials in each period. '
        'Per-period denominators (n) appear below the x-axis. The 2018-2019 '
        'period is the lowest-reporting bin (30%, with 0 of 4 trials in 2018 '
        'reporting race/ethnicity).')
    doc.add_page_break()

    # ---- Table 2 ----
    add_apa_title_block(doc, 'Table 2',
        'PROGRESS-Plus Reporting Rates With 95% Confidence Intervals (N = 63)')
    t2 = generate_table2_data()
    table2 = doc.add_table(rows=len(t2) + 1, cols=4)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['Variable', 'n/N', '%', '95% CI']):
        format_table_cell(table2.rows[0].cells[j], h, bold=True,
                          align='center' if j > 0 else 'left')
    for i, row_data in enumerate(t2):
        is_section = row_data[1] == '' and row_data[2] == ''
        for j, val in enumerate(row_data):
            align = 'center' if j > 0 else 'left'
            format_table_cell(table2.rows[i + 1].cells[j], val,
                              bold=is_section, align=align)
    apply_apa_table_borders(table2)
    add_table_note(doc,
        '95% CI computed using Wilson score interval. \u2014 indicates CI '
        'not applicable for sub-category breakdowns. Race/ethnicity '
        'sub-rows use the 36-study reporter denominator. SOGI = sexual '
        'orientation and gender identity.')
    doc.add_page_break()

    # ---- Figure 4 ----
    add_apa_title_block(doc, 'Figure 4',
                        'Race / Ethnicity Reporting by Funder Type')
    doc.add_picture(os.path.join(FIG_DIR, 'figure4_race_by_funder.png'),
                    width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table_note(doc,
        'Stacked horizontal bars show the number of trials reporting versus '
        'not reporting race/ethnicity within each funder category. The '
        'right-side label gives the percent of each funder group that '
        'reported race/ethnicity. CIHR-funded trials had the highest absolute '
        'count of non-reporting (5/16, 31%).')
    doc.add_page_break()

    # ---- Table 3 ----
    add_apa_title_block(doc, 'Table 3',
        'Race / Ethnicity Reporting by Funder Type: Cross-Tabulation')
    t3, t3_p = generate_table3_data()
    table3 = doc.add_table(rows=len(t3) + 1, cols=6)
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['Funder Type', 'Yes', 'Partial', 'No',
                           'Total', '% Reported']):
        format_table_cell(table3.rows[0].cells[j], h, bold=True,
                          align='center' if j > 0 else 'left')
    for i, row_data in enumerate(t3):
        for j, val in enumerate(row_data):
            align = 'center' if j > 0 else 'left'
            format_table_cell(table3.rows[i + 1].cells[j], val, align=align)
    apply_apa_table_borders(table3)
    add_table_note(doc, f'{t3_p}. "% Reported" counts strict "Yes" responses '
                        'only. CIHR = Canadian Institutes of Health Research. '
                        'NIH = National Institutes of Health.')
    doc.add_page_break()

    # ---- Figure 5 ----
    add_apa_title_block(doc, 'Figure 5',
        'Indigenous Participation and Governance Indicators (7 Trials)')
    doc.add_picture(os.path.join(FIG_DIR, 'figure5_indigenous_heatmap.png'),
                    width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table_note(doc,
        'Filled circles indicate Yes; gray circles indicate partial or '
        'non-specific reporting; open circles indicate No. The vertical '
        'rule separates "Documented" from "Governance & OCAP" indicators. '
        'OCAP = Ownership, Control, Access, Possession (First Nations '
        'principles). No trial referenced OCAP principles or any Indigenous '
        'data governance framework.')
    doc.add_page_break()

    # ---- Table 4 ----
    add_apa_title_block(doc, 'Table 4',
        'Detailed Indigenous Participation Characteristics (7 Trials)')
    t4 = generate_table4_data()
    table4 = doc.add_table(rows=len(t4) + 1, cols=8)
    table4.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(['Rec #', 'Study', 'Disorder', 'N', 'Groups',
                           'Partnership', 'Sov.', 'Notes']):
        format_table_cell(table4.rows[0].cells[j], h, bold=True, size=9,
                          align='center' if j in (0, 3, 6) else 'left')
    for i, row_data in enumerate(t4):
        for j, val in enumerate(row_data):
            align = 'center' if j in (0, 3, 6) else 'left'
            format_table_cell(table4.rows[i + 1].cells[j], val, size=9,
                              align=align)
    apply_apa_table_borders(table4)
    add_table_note(doc,
        'Sov. = Indigenous Data Sovereignty Score (0 = no mention; '
        '1 = statement only; 2 = protocol referenced; 3 = full governance). '
        'All seven trials scored 0. OCAP = Ownership, Control, Access, '
        'Possession.')
    doc.add_page_break()

    # ---- Figure 6 ----
    add_apa_title_block(doc, 'Figure 6', 'Sex vs. Gender Reporting (N = 63)')
    doc.add_picture(os.path.join(FIG_DIR, 'figure6_sex_gender.png'),
                    width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table_note(doc,
        'Stacked horizontal bar showing the four mutually exclusive '
        'categories. All four trials that distinguished sex from gender '
        'were published 2022\u20132025. "Gender only" comprises three '
        'trials reporting gender identity but not biological sex.')
    doc.add_page_break()

    # ---- Figure 7 ----
    add_apa_title_block(doc, 'Figure 7',
                        'CONSORT-Equity Compliance (N = 63)')
    doc.add_picture(os.path.join(FIG_DIR, 'figure7_consort_equity.png'),
                    width=Inches(6.0))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table_note(doc,
        'CONSORT-Equity compliance assessed against Welch et al. (2017): '
        'sociodemographic characteristics reported, equity-relevant '
        'subgroup analyses, and generalizability discussion with explicit '
        'equity reference. No trial achieved full compliance on all three '
        'criteria (0 / 63, 0%).')
    doc.add_page_break()

    # ---- Figure 8 ----
    add_apa_title_block(doc, 'Figure 8',
        'PROGRESS-Plus Reporting Completeness by Publication Year and '
        'Sample Size')
    doc.add_picture(os.path.join(FIG_DIR, 'figure8_bubble_completeness.png'),
                    width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_table_note(doc,
        'Completeness score = count of PROGRESS-Plus variables reported '
        '(0\u201313). Marker area is proportional to total sample size. '
        'Marker shape indicates disorder category. The dashed line is the '
        'ordinary least-squares trend (slope and significance shown in the '
        'legend).')

    doc.save(DOCX_PATH)
    print(f"\n  Word document saved: {DOCX_PATH}")


# =============================================================================
# VALIDATION
# =============================================================================
def run_validation():
    print("\n=== VALIDATION ===")
    assert N == 63, f"FAIL: N={N}, expected 63"
    print(f"  Denominator check: PASS (N={N})")

    n_race = (df['race_reported'].str.strip().str.lower() == 'yes').sum()
    checks = {
        'Race reported (Yes)':       (n_race, 36),
        'Sex reported':              (count_reported(df['sex_reported']), 60),
        'Gender reported':           (count_reported(df['gender_reported']), 9),
        'Sex/gender distinguished':  ((df['sex_gender_distinguished']
                                       .str.strip().str.lower() == 'yes').sum(), 4),
        'SOGI':                      (count_reported(df['sogi_reported']), 2),
        'Religion':                  (count_reported(df['religion_reported']), 0),
        'Intersectional':            (count_reported(df['intersectional_analysis']), 0),
        'Indigenous participation':  ((df['indigenous_participation'] == 'Yes').sum(), 7),
        'CONSORT full':              ((df['consort_equity_compliant']
                                       .str.strip().str.lower() == 'yes').sum(), 0),
        'CONSORT partial':           ((df['consort_equity_compliant']
                                       .str.strip().str.lower() == 'partial').sum(), 17),
        'CONSORT none':              ((df['consort_equity_compliant']
                                       .str.strip().str.lower() == 'no').sum(), 46),
    }
    all_pass = True
    for name, (actual, expected) in checks.items():
        status = "PASS" if actual == expected else "DISCREPANCY"
        if status == "DISCREPANCY":
            all_pass = False
        print(f"  {name}: {actual} (expected {expected}) -- {status}")

    for f in ['figure1_prisma_flow.png', 'figure2_progress_plus_bars.png',
              'figure3_race_over_time.png', 'figure4_race_by_funder.png',
              'figure5_indigenous_heatmap.png', 'figure6_sex_gender.png',
              'figure7_consort_equity.png', 'figure8_bubble_completeness.png']:
        path = os.path.join(FIG_DIR, f)
        exists = os.path.exists(path)
        size = os.path.getsize(path) if exists else 0
        print(f"  {f}: {'OK' if exists else 'MISSING'} ({size:,} bytes)")

    docx_exists = os.path.exists(DOCX_PATH)
    docx_size = os.path.getsize(DOCX_PATH) if docx_exists else 0
    print(f"  tables_and_figures.docx: "
          f"{'OK' if docx_exists else 'MISSING'} ({docx_size:,} bytes)")
    return all_pass


# =============================================================================
# MAIN
# =============================================================================
if __name__ == '__main__':
    print("Generating figures and tables for EDIA scoping review...")
    print(f"Data: {N} studies (records #4, #12, #45 excluded)\n")

    print("Generating figures:")
    generate_figure1()
    generate_figure2()
    generate_figure3()
    generate_figure4()
    generate_figure5()
    generate_figure6()
    generate_figure7()
    generate_figure8()

    print("\nBuilding Word document:")
    build_docx()

    ok = run_validation()
    print(f"\n{'=' * 50}")
    print(f"COMPLETE. Validation: {'PASS' if ok else 'REVIEW DISCREPANCIES'}")
    print(f"{'=' * 50}")
